from __future__ import annotations

from dataclasses import dataclass
from datetime import date, datetime
from decimal import Decimal
from io import BytesIO, StringIO
from typing import List, Optional, Union

from app.schemas.breakeven_report import BreakevenSummaryReportData, BreakevenSummaryRow
from app.schemas.report import (
    AnimalProjectionBlock,
    FrozenContainersMonthlyReportData,
    FrozenContainersReportRow,
    OrdersMonthlyReportData,
    OrdersReportRow,
    OrdersReportSection,
)


@dataclass
class GeneratedReport:
    filename: str
    media_type: str
    content: bytes


CSV_MEDIA_TYPE = "text/csv"
PDF_MEDIA_TYPE = "application/pdf"
DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

ReportData = Union[
    OrdersMonthlyReportData,
    FrozenContainersMonthlyReportData,
    BreakevenSummaryReportData,
]


# =============================================================================
# Filename generation
# =============================================================================
def _build_report_filename(report_type: str, report_format: str) -> str:
    friendly_names = {
        "orders_monthly": "Order Confirmed Report",
        "frozen_containers_monthly": "Frozen Confirmed Containers",
        "breakeven_summary": "Breakeven Summary Report",
    }
    base_name = friendly_names.get(
        (report_type or "").strip().lower(),
        (report_type or "Report").replace("_", " ").title(),
    )
    today_str = datetime.now().strftime("%Y%m%d")
    safe_format = (report_format or "pdf").strip().lower()
    return f"{base_name} {today_str}.{safe_format}"


# =============================================================================
# Formatting helpers
# =============================================================================
def _format_decimal(value: Optional[Decimal | int | float | str]) -> str:
    if value is None:
        return "0.00"
    if isinstance(value, Decimal):
        return f"{value:,.2f}"
    try:
        return f"{Decimal(str(value)):,.2f}"
    except Exception:
        return str(value)


def _format_decimal4(value: Optional[Decimal | int | float | str]) -> str:
    if value is None:
        return "0.0000"
    if isinstance(value, Decimal):
        return f"{value:,.4f}"
    try:
        return f"{Decimal(str(value)):,.4f}"
    except Exception:
        return str(value)


def _format_int(value: Optional[int]) -> str:
    return str(int(value or 0))


def _format_currency(value: Optional[Decimal | int | float | str]) -> str:
    return f"${_format_decimal(value)}"


def _format_human_date(value: Optional[date]) -> str:
    if not value:
        return "-"
    return value.strftime("%d-%b-%y")


def _format_compact_date(value: Optional[datetime]) -> str:
    if not value:
        return ""
    return value.strftime("%d.%m.%Y")


def _format_report_day(value: Optional[date | datetime]) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%d.%m.%Y")
    return value.strftime("%d.%m.%Y")


def _safe_text(value: Optional[str]) -> str:
    return (value or "").strip()


def _status_text(value: str) -> str:
    return (value or "").replace("_", " ").title()


def _projection_title(report_data: OrdersMonthlyReportData) -> str:
    if report_data.animal_projection and _safe_text(report_data.animal_projection.title):
        return report_data.animal_projection.title
    try:
        dt = date(report_data.year, report_data.month, 1)
        return f"ANIMAL REQUIREMENTS PROJECTION FOR THE MONTH OF {dt.strftime('%B').upper()}"
    except Exception:
        return "ANIMAL REQUIREMENTS PROJECTION"


def _summary_title(report_data: OrdersMonthlyReportData) -> str:
    return _safe_text(report_data.subtitle) or "MONTHLY SUMMARY REPORT"


# =============================================================================
# Projection helpers - force 4 weeks for every month
# =============================================================================
def _week_index_from_label(label: str) -> Optional[int]:
    text = (label or "").strip().lower()
    if text in {"1st week", "week 1", "1"}:
        return 1
    if text in {"2nd week", "week 2", "2"}:
        return 2
    if text in {"3rd week", "week 3", "3"}:
        return 3
    if text in {"4th week", "week 4", "4"}:
        return 4
    return None


def _week_label_from_index(index: int) -> str:
    mapping = {
        1: "1st Week",
        2: "2nd Week",
        3: "3rd Week",
        4: "4th Week",
    }
    return mapping.get(index, f"Week {index}")


def _normalize_projection_rows(
    projection: Optional[AnimalProjectionBlock],
) -> list[dict[str, int | str]]:
    buckets: dict[int, dict[str, int | str]] = {
        1: {"label": "1st Week", "goats": 0, "sheep": 0, "cattle": 0, "total_animals": 0},
        2: {"label": "2nd Week", "goats": 0, "sheep": 0, "cattle": 0, "total_animals": 0},
        3: {"label": "3rd Week", "goats": 0, "sheep": 0, "cattle": 0, "total_animals": 0},
        4: {"label": "4th Week", "goats": 0, "sheep": 0, "cattle": 0, "total_animals": 0},
    }

    if projection and projection.rows:
        for row in projection.rows:
            week_idx = _week_index_from_label(getattr(row, "label", "") or "")
            if week_idx not in {1, 2, 3, 4}:
                continue

            goats = int(getattr(row, "goats", 0) or 0)
            sheep = int(getattr(row, "sheep", 0) or 0)
            cattle = int(getattr(row, "cattle", 0) or 0)

            buckets[week_idx] = {
                "label": _week_label_from_index(week_idx),
                "goats": goats,
                "sheep": sheep,
                "cattle": cattle,
                "total_animals": goats + sheep + cattle,
            }

    return [buckets[1], buckets[2], buckets[3], buckets[4]]


def _projection_totals_from_normalized_rows(
    normalized_rows: list[dict[str, int | str]],
) -> tuple[int, int, int, int]:
    total_goats = sum(int(row["goats"]) for row in normalized_rows)
    total_sheep = sum(int(row["sheep"]) for row in normalized_rows)
    total_cattle = sum(int(row["cattle"]) for row in normalized_rows)
    grand_total = sum(int(row["total_animals"]) for row in normalized_rows)
    return total_goats, total_sheep, total_cattle, grand_total


# =============================================================================
# Orders report text helpers
# =============================================================================
def _preferred_species_label(summary_text: str, animal_key: str) -> str:
    text = (summary_text or "").lower()

    if animal_key == "goat":
        return "Goat"

    if animal_key == "sheep":
        if "mutton" in text:
            return "Mutton"
        if "lamb" in text:
            return "Lamb"
        return "Sheep"

    if animal_key == "cattle":
        if "beef" in text:
            return "Beef"
        return "Cattle"

    return animal_key.title()


def _pieces_text_for_row(row: OrdersReportRow) -> str:
    parts: list[str] = []

    if int(row.goat_pieces_required or 0) > 0:
        parts.append(
            f"{_preferred_species_label(row.product_summary or '', 'goat')} {int(row.goat_pieces_required)} pcs"
        )

    if int(row.sheep_pieces_required or 0) > 0:
        parts.append(
            f"{_preferred_species_label(row.product_summary or '', 'sheep')} {int(row.sheep_pieces_required)} pcs"
        )

    if int(row.cattle_pieces_required or 0) > 0:
        parts.append(
            f"{_preferred_species_label(row.product_summary or '', 'cattle')} {int(row.cattle_pieces_required)} pcs"
        )

    if not parts:
        return f"{_format_int(row.total_pieces_required)} pcs"

    return ", ".join(parts)


def _quantity_text_for_row(row: OrdersReportRow) -> str:
    return row.product_summary or f"{_format_decimal(row.total_quantity_kg)} kg"


def _build_orders_summary_items(
    report_data: OrdersMonthlyReportData,
) -> list[tuple[str, str, str]]:
    items: list[tuple[str, str, str]] = []

    if not report_data.summary:
        return items

    items.append(("1", "Total orders per Month(kgs)", _format_decimal(report_data.summary.total_quantity_kg)))
    items.append(("2", "Breakeven point(kgs)", _format_decimal(report_data.summary.breakeven_quantity_kg)))
    items.append(
        ("3", "Total breakeven quantities achieved (%)", _format_decimal(report_data.summary.breakeven_achieved_percentage))
    )
    items.append(
        ("4", "Balance to hit breakeven quantities (Kgs)", _format_decimal(report_data.summary.breakeven_balance_quantity_kg))
    )
    items.append(("5", "Total pieces required", _format_int(report_data.summary.total_pieces_required)))
    items.append(("6", "Total animals required", _format_int(report_data.summary.total_animals_required)))
    items.append(("7", "Number of orders", _format_int(report_data.summary.total_orders)))

    return items


def _should_show_overall_totals(report_data: OrdersMonthlyReportData) -> bool:
    if not report_data.totals:
        return False
    return len(report_data.sections) > 1


# =============================================================================
# Breakeven helpers
# =============================================================================
def _build_breakeven_rows(report_data: BreakevenSummaryReportData) -> list[BreakevenSummaryRow]:
    return list(report_data.rows or [])


# =============================================================================
# CSV generation
# =============================================================================
def _build_csv_for_orders_monthly(report_data: OrdersMonthlyReportData) -> bytes:
    import csv

    output = StringIO()
    writer = csv.writer(output)

    writer.writerow([f"{report_data.organization_name or ''} – {report_data.title}"])
    writer.writerow(["Prepared by", report_data.prepared_by or ""])
    writer.writerow(["Date", _format_compact_date(report_data.generated_at)])
    writer.writerow([_summary_title(report_data)])
    writer.writerow([])

    if report_data.summary:
        writer.writerow([_summary_title(report_data)])
        writer.writerow(["No", "Particulars", "Value"])
        for no, label, value in _build_orders_summary_items(report_data):
            writer.writerow([no, label, value])
        writer.writerow([])

    if report_data.animal_projection:
        normalized_rows = _normalize_projection_rows(report_data.animal_projection)
        total_goats, total_sheep, total_cattle, grand_total = _projection_totals_from_normalized_rows(normalized_rows)

        writer.writerow([_projection_title(report_data)])
        writer.writerow(["WEEK", "GOAT", "SHEEP", "CATTLE", "TOTAL"])

        for row in normalized_rows:
            writer.writerow(
                [
                    row["label"],
                    row["goats"],
                    row["sheep"],
                    row["cattle"],
                    row["total_animals"],
                ]
            )

        writer.writerow(
            [
                "TOTAL (All Animals)",
                total_goats,
                total_sheep,
                total_cattle,
                grand_total,
            ]
        )
        writer.writerow([])

    for section in report_data.sections:
        writer.writerow([section.section_title])
        writer.writerow(
            [
                "Serial No.",
                "Name of Enterprise",
                "Product Quantity",
                "Pieces Required",
                "Slaughter Schedule",
                "Expected Delivery",
                "Status",
            ]
        )

        for row in section.rows:
            writer.writerow(
                [
                    f"{row.serial_no}.",
                    row.enterprise_name,
                    _quantity_text_for_row(row),
                    _pieces_text_for_row(row),
                    _format_human_date(row.slaughter_schedule),
                    _format_human_date(row.expected_delivery),
                    _status_text(row.status),
                ]
            )

        writer.writerow(
            [
                "Totals",
                "",
                f"{_format_decimal(section.total_quantity_kg)} kg",
                f"{section.total_pieces_required:,} pcs",
                "",
                "",
                "",
            ]
        )
        writer.writerow([])

    if _should_show_overall_totals(report_data):
        writer.writerow(
            [
                "OVERALL TOTALS",
                "",
                f"{_format_decimal(report_data.totals.total_quantity_kg)} kg",
                f"{report_data.totals.total_pieces_required:,} pcs",
                "",
                "",
                "",
            ]
        )

    return output.getvalue().encode("utf-8")


def _build_csv_for_frozen_containers_monthly(
    report_data: FrozenContainersMonthlyReportData,
) -> bytes:
    import csv

    output = StringIO()
    writer = csv.writer(output)

    writer.writerow([f"{report_data.organization_name or ''} – {report_data.title}"])
    writer.writerow(["Prepared by", report_data.prepared_by or ""])
    writer.writerow(["Date", _format_compact_date(report_data.generated_at)])
    writer.writerow([])

    writer.writerow(
        [
            "No.",
            "Client Name",
            "Order Ratio",
            "Status",
            "Container Value (USD)",
            "Price per kg",
            "Down Payment",
            "Balance",
            "Container Gate in",
            "Departure Date",
            "Jurisdiction",
        ]
    )

    for row in report_data.rows:
        writer.writerow(
            [
                row.serial_no,
                row.client_name,
                row.order_ratio or "",
                _status_text(row.status),
                _format_currency(row.container_value_usd) if row.container_value_usd is not None else "",
                _format_decimal4(row.price_per_kg_usd) if row.price_per_kg_usd is not None else "",
                _format_currency(row.down_payment_usd) if row.down_payment_usd is not None else "",
                _format_currency(row.balance_usd) if row.balance_usd is not None else "",
                _format_human_date(row.container_gate_in),
                _format_human_date(row.departure_date),
                row.jurisdiction or "",
            ]
        )

    if report_data.totals:
        writer.writerow([])
        writer.writerow(
            [
                "Totals",
                "",
                "",
                "",
                _format_currency(report_data.totals.total_container_value_usd),
                "",
                _format_currency(report_data.totals.total_down_payment_usd),
                _format_currency(report_data.totals.total_balance_usd),
                "",
                "",
                "",
            ]
        )

    return output.getvalue().encode("utf-8")


def _build_csv_for_breakeven_summary(report_data: BreakevenSummaryReportData) -> bytes:
    import csv

    output = StringIO()
    writer = csv.writer(output)

    writer.writerow([f"{report_data.organization_name or ''} – {report_data.title}"])
    writer.writerow(["Prepared by", report_data.prepared_by or ""])
    writer.writerow(["Generated", _format_compact_date(report_data.generated_at)])
    writer.writerow(["Report Date", _format_report_day(report_data.report_date)])
    writer.writerow(["Month", report_data.month or ""])
    writer.writerow(["Year", report_data.year or ""])
    writer.writerow([])

    writer.writerow(["Index", "Metric", "Quantity (Tonnes)", "USD (Total)", "Percentage"])

    for row in _build_breakeven_rows(report_data):
        writer.writerow(
            [
                row.index,
                row.metric,
                row.quantity_display or "",
                row.usd_display or "",
                row.percentage_display or "",
            ]
        )

    return output.getvalue().encode("utf-8")


# =============================================================================
# PDF generation
# =============================================================================
def _require_reportlab():
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import (
            PageBreak,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )
    except ImportError as exc:
        raise RuntimeError(
            "PDF generation requires reportlab. Install it with: pip install reportlab"
        ) from exc

    return {
        "colors": colors,
        "TA_CENTER": TA_CENTER,
        "TA_LEFT": TA_LEFT,
        "A4": A4,
        "landscape": landscape,
        "ParagraphStyle": ParagraphStyle,
        "getSampleStyleSheet": getSampleStyleSheet,
        "mm": mm,
        "PageBreak": PageBreak,
        "Paragraph": Paragraph,
        "SimpleDocTemplate": SimpleDocTemplate,
        "Spacer": Spacer,
        "Table": Table,
        "TableStyle": TableStyle,
    }


def _pdf_page_decor(canvas, doc):
    from reportlab.lib import colors

    page_width, page_height = doc.pagesize

    green = colors.HexColor("#2F8F2F")
    red = colors.HexColor("#8B0000")
    muted = colors.HexColor("#666666")

    canvas.saveState()

    canvas.setFillColor(green)
    canvas.rect(doc.leftMargin - 3, page_height - 16, 18, 4, fill=1, stroke=0)

    canvas.setFillColor(red)
    canvas.rect(
        doc.leftMargin + 21,
        page_height - 16,
        page_width - doc.leftMargin - doc.rightMargin - 21,
        4,
        fill=1,
        stroke=0,
    )

    canvas.setFont("Times-Roman", 8)
    canvas.setFillColor(muted)
    canvas.drawRightString(page_width - doc.rightMargin, 12, f"Page {canvas.getPageNumber()}")

    canvas.restoreState()


def _build_pdf_styles(ParagraphStyle, getSampleStyleSheet, TA_LEFT, TA_CENTER):
    styles = getSampleStyleSheet()

    return {
        "title": ParagraphStyle(
            "ReportTitle",
            parent=styles["Title"],
            fontName="Times-Bold",
            fontSize=18,
            leading=22,
            alignment=TA_LEFT,
            spaceAfter=4,
        ),
        "meta": ParagraphStyle(
            "ReportMeta",
            parent=styles["BodyText"],
            fontName="Times-Bold",
            fontSize=10,
            leading=12,
            alignment=TA_LEFT,
            spaceAfter=0,
        ),
        "mini_heading": ParagraphStyle(
            "MiniHeading",
            parent=styles["Heading3"],
            fontName="Times-Bold",
            fontSize=12,
            leading=14,
            alignment=TA_CENTER,
            spaceBefore=3,
            spaceAfter=6,
        ),
        "section": ParagraphStyle(
            "ReportSection",
            parent=styles["Heading3"],
            fontName="Times-Bold",
            fontSize=12,
            leading=14,
            alignment=TA_CENTER,
            spaceBefore=4,
            spaceAfter=6,
        ),
        "cell": ParagraphStyle(
            "Cell",
            parent=styles["BodyText"],
            fontName="Times-Roman",
            fontSize=10,
            leading=11.5,
            alignment=TA_LEFT,
        ),
        "cell_bold": ParagraphStyle(
            "CellBold",
            parent=styles["BodyText"],
            fontName="Times-Bold",
            fontSize=10,
            leading=11.5,
            alignment=TA_LEFT,
        ),
        "header_cell": ParagraphStyle(
            "HeaderCell",
            parent=styles["BodyText"],
            fontName="Times-Bold",
            fontSize=9,
            leading=10.5,
            alignment=TA_LEFT,
        ),
        "summary_no": ParagraphStyle(
            "SummaryNo",
            parent=styles["BodyText"],
            fontName="Times-Bold",
            fontSize=10,
            leading=12,
            alignment=TA_LEFT,
        ),
        "summary_value": ParagraphStyle(
            "SummaryValue",
            parent=styles["BodyText"],
            fontName="Times-Roman",
            fontSize=10,
            leading=12,
            alignment=TA_LEFT,
        ),
    }


def _pdf_p(text: str, Paragraph, style):
    return Paragraph(_safe_text(text).replace("\n", "<br/>"), style)


def _build_pdf_meta_table(report_data: ReportData, Table, TableStyle):
    prepared = getattr(report_data, "prepared_by", "") or ""
    generated = _format_compact_date(getattr(report_data, "generated_at", None))
    data = [[f"Prepared by: {prepared}", f"Date: {generated}"]]

    table = Table(data, colWidths=[210, 120])
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), "Times-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    return table


def _build_pdf_summary_table(summary_items, Table, TableStyle, Paragraph, colors, styles):
    data = [
        [
            _pdf_p("No", Paragraph, styles["header_cell"]),
            _pdf_p("Particulars", Paragraph, styles["header_cell"]),
            _pdf_p("Value", Paragraph, styles["header_cell"]),
        ]
    ]

    for no, label, value in summary_items:
        data.append(
            [
                _pdf_p(no, Paragraph, styles["summary_no"]),
                _pdf_p(label, Paragraph, styles["cell"]),
                _pdf_p(value, Paragraph, styles["summary_value"]),
            ]
        )

    table = Table(data, colWidths=[24, 235, 95])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EFEFEF")),
                ("GRID", (0, 0), (-1, -1), 0.6, colors.black),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("PADDING", (0, 0), (-1, -1), 5),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    return table


def _build_pdf_projection_table(report_data, projection, Table, TableStyle, Paragraph, colors, styles):
    normalized_rows = _normalize_projection_rows(projection)
    total_goats, total_sheep, total_cattle, grand_total = _projection_totals_from_normalized_rows(normalized_rows)

    data = [
        [
            _pdf_p("WEEK", Paragraph, styles["header_cell"]),
            _pdf_p("GOAT", Paragraph, styles["header_cell"]),
            _pdf_p("SHEEP", Paragraph, styles["header_cell"]),
            _pdf_p("CATTLE", Paragraph, styles["header_cell"]),
            _pdf_p("TOTAL", Paragraph, styles["header_cell"]),
        ]
    ]

    for row in normalized_rows:
        data.append(
            [
                _pdf_p(str(row["label"]), Paragraph, styles["cell"]),
                _pdf_p(str(row["goats"]), Paragraph, styles["summary_value"]),
                _pdf_p(str(row["sheep"]), Paragraph, styles["summary_value"]),
                _pdf_p(str(row["cattle"]), Paragraph, styles["summary_value"]),
                _pdf_p(str(row["total_animals"]), Paragraph, styles["summary_value"]),
            ]
        )

    data.append(
        [
            _pdf_p("TOTAL (All Animals)", Paragraph, styles["cell_bold"]),
            _pdf_p(str(total_goats), Paragraph, styles["cell_bold"]),
            _pdf_p(str(total_sheep), Paragraph, styles["cell_bold"]),
            _pdf_p(str(total_cattle), Paragraph, styles["cell_bold"]),
            _pdf_p(str(grand_total), Paragraph, styles["cell_bold"]),
        ]
    )

    table = Table(data, colWidths=[170, 70, 70, 70, 75])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EFEFEF")),
                ("GRID", (0, 0), (-1, -1), 0.6, colors.black),
                ("TEXTCOLOR", (0, len(data) - 1), (-1, len(data) - 1), colors.HexColor("#B22222")),
                ("FONTNAME", (0, len(data) - 1), (-1, len(data) - 1), "Times-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("PADDING", (0, 0), (-1, -1), 5),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    return table


def _build_pdf_page_one_orders(
    report_data: OrdersMonthlyReportData,
    story,
    Table,
    TableStyle,
    Paragraph,
    Spacer,
    colors,
    styles,
):
    story.append(_pdf_p(f"{report_data.organization_name or ''} – {report_data.title}", Paragraph, styles["title"]))
    story.append(Spacer(1, 2))
    story.append(_build_pdf_meta_table(report_data, Table, TableStyle))
    story.append(Spacer(1, 8))

    if report_data.summary:
        story.append(_pdf_p(_summary_title(report_data), Paragraph, styles["mini_heading"]))
        story.append(
            _build_pdf_summary_table(
                _build_orders_summary_items(report_data),
                Table,
                TableStyle,
                Paragraph,
                colors,
                styles,
            )
        )
        story.append(Spacer(1, 10))

    if report_data.animal_projection is not None:
        story.append(_pdf_p(_projection_title(report_data), Paragraph, styles["mini_heading"]))
        story.append(
            _build_pdf_projection_table(
                report_data,
                report_data.animal_projection,
                Table,
                TableStyle,
                Paragraph,
                colors,
                styles,
            )
        )


def _build_pdf_page_one_frozen(
    report_data: FrozenContainersMonthlyReportData,
    story,
    Table,
    TableStyle,
    Paragraph,
    Spacer,
    styles,
):
    story.append(_pdf_p(f"{report_data.organization_name or ''} – {report_data.title}", Paragraph, styles["title"]))
    story.append(Spacer(1, 2))
    story.append(_build_pdf_meta_table(report_data, Table, TableStyle))
    story.append(Spacer(1, 8))


def _build_pdf_page_one_breakeven(
    report_data: BreakevenSummaryReportData,
    story,
    Table,
    TableStyle,
    Paragraph,
    Spacer,
    styles,
):
    story.append(_pdf_p(f"{report_data.organization_name or ''} – {report_data.title}", Paragraph, styles["title"]))
    story.append(Spacer(1, 2))
    story.append(_build_pdf_meta_table(report_data, Table, TableStyle))
    story.append(Spacer(1, 6))
    story.append(_pdf_p(f"Report Date: {_format_report_day(report_data.report_date)}", Paragraph, styles["meta"]))
    story.append(_pdf_p(f"Month/Year: {report_data.month or ''}/{report_data.year or ''}", Paragraph, styles["meta"]))
    story.append(Spacer(1, 8))


def _build_pdf_orders_table(section, Table, TableStyle, Paragraph, colors, styles):
    data = [
        [
            _pdf_p("Serial No.", Paragraph, styles["header_cell"]),
            _pdf_p("Name of Enterprise", Paragraph, styles["header_cell"]),
            _pdf_p("Product Quantity", Paragraph, styles["header_cell"]),
            _pdf_p("Pieces Required", Paragraph, styles["header_cell"]),
            _pdf_p("Slaughter Schedule", Paragraph, styles["header_cell"]),
            _pdf_p("Expected Delivery", Paragraph, styles["header_cell"]),
            _pdf_p("Status", Paragraph, styles["header_cell"]),
        ]
    ]

    for row in section.rows:
        data.append(
            [
                _pdf_p(f"{row.serial_no}.", Paragraph, styles["summary_value"]),
                _pdf_p(row.enterprise_name, Paragraph, styles["cell"]),
                _pdf_p(_quantity_text_for_row(row), Paragraph, styles["cell"]),
                _pdf_p(_pieces_text_for_row(row), Paragraph, styles["cell"]),
                _pdf_p(_format_human_date(row.slaughter_schedule), Paragraph, styles["summary_value"]),
                _pdf_p(_format_human_date(row.expected_delivery), Paragraph, styles["summary_value"]),
                _pdf_p(_status_text(row.status), Paragraph, styles["summary_value"]),
            ]
        )

    data.append(
        [
            _pdf_p("Totals", Paragraph, styles["cell_bold"]),
            "",
            _pdf_p(f"{_format_decimal(section.total_quantity_kg)} kg", Paragraph, styles["cell_bold"]),
            _pdf_p(f"{section.total_pieces_required:,} pcs", Paragraph, styles["cell_bold"]),
            "",
            "",
            "",
        ]
    )

    table = Table(
        data,
        repeatRows=1,
        colWidths=[58, 118, 180, 128, 78, 78, 68],
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EFEFEF")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def _build_pdf_frozen_table(
    rows: List[FrozenContainersReportRow],
    totals,
    Table,
    TableStyle,
    Paragraph,
    colors,
    styles,
):
    data = [
        [
            _pdf_p("No.", Paragraph, styles["header_cell"]),
            _pdf_p("Client Name", Paragraph, styles["header_cell"]),
            _pdf_p("Order Ratio", Paragraph, styles["header_cell"]),
            _pdf_p("Status", Paragraph, styles["header_cell"]),
            _pdf_p("Container Value (USD)", Paragraph, styles["header_cell"]),
            _pdf_p("Price per kg", Paragraph, styles["header_cell"]),
            _pdf_p("Down Payment", Paragraph, styles["header_cell"]),
            _pdf_p("Balance", Paragraph, styles["header_cell"]),
            _pdf_p("Container Gate in", Paragraph, styles["header_cell"]),
            _pdf_p("Departure Date", Paragraph, styles["header_cell"]),
            _pdf_p("Jurisdiction", Paragraph, styles["header_cell"]),
        ]
    ]

    for row in rows:
        data.append(
            [
                _pdf_p(str(row.serial_no), Paragraph, styles["summary_value"]),
                _pdf_p(row.client_name, Paragraph, styles["cell"]),
                _pdf_p(row.order_ratio or "", Paragraph, styles["cell"]),
                _pdf_p(_status_text(row.status), Paragraph, styles["summary_value"]),
                _pdf_p(_format_currency(row.container_value_usd) if row.container_value_usd is not None else "-", Paragraph, styles["summary_value"]),
                _pdf_p(_format_decimal4(row.price_per_kg_usd) if row.price_per_kg_usd is not None else "-", Paragraph, styles["summary_value"]),
                _pdf_p(_format_currency(row.down_payment_usd) if row.down_payment_usd is not None else "-", Paragraph, styles["summary_value"]),
                _pdf_p(_format_currency(row.balance_usd) if row.balance_usd is not None else "-", Paragraph, styles["summary_value"]),
                _pdf_p(_format_human_date(row.container_gate_in), Paragraph, styles["summary_value"]),
                _pdf_p(_format_human_date(row.departure_date), Paragraph, styles["summary_value"]),
                _pdf_p(row.jurisdiction or "", Paragraph, styles["cell"]),
            ]
        )

    has_totals = totals is not None
    if has_totals:
        data.append(
            [
                _pdf_p("Totals", Paragraph, styles["cell_bold"]),
                "",
                "",
                "",
                _pdf_p(_format_currency(totals.total_container_value_usd), Paragraph, styles["cell_bold"]),
                "",
                _pdf_p(_format_currency(totals.total_down_payment_usd), Paragraph, styles["cell_bold"]),
                _pdf_p(_format_currency(totals.total_balance_usd), Paragraph, styles["cell_bold"]),
                "",
                "",
                "",
            ]
        )

    table = Table(
        data,
        repeatRows=1,
        colWidths=[44, 98, 64, 58, 84, 56, 72, 72, 64, 64, 70],
    )

    style_commands = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EFEFEF")),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]

    if has_totals:
        last_row = len(data) - 1
        style_commands.extend(
            [
                ("SPAN", (0, last_row), (3, last_row)),
                ("FONTNAME", (0, last_row), (-1, last_row), "Times-Bold"),
                ("BACKGROUND", (0, last_row), (-1, last_row), colors.HexColor("#F7F7F7")),
            ]
        )

    table.setStyle(TableStyle(style_commands))
    return table


def _build_pdf_breakeven_table(
    rows: List[BreakevenSummaryRow],
    Table,
    TableStyle,
    Paragraph,
    colors,
    styles,
):
    data = [
        [
            _pdf_p("Index", Paragraph, styles["header_cell"]),
            _pdf_p("Metric", Paragraph, styles["header_cell"]),
            _pdf_p("Quantity (Tonnes)", Paragraph, styles["header_cell"]),
            _pdf_p("USD (Total)", Paragraph, styles["header_cell"]),
            _pdf_p("Percentage", Paragraph, styles["header_cell"]),
        ]
    ]

    for row in rows:
        data.append(
            [
                _pdf_p(str(row.index), Paragraph, styles["summary_value"]),
                _pdf_p(row.metric, Paragraph, styles["cell"]),
                _pdf_p(row.quantity_display or "", Paragraph, styles["summary_value"]),
                _pdf_p(row.usd_display or "", Paragraph, styles["summary_value"]),
                _pdf_p(row.percentage_display or "", Paragraph, styles["summary_value"]),
            ]
        )

    table = Table(
        data,
        repeatRows=1,
        colWidths=[42, 230, 100, 100, 82],
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EFEFEF")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def _build_pdf_for_orders_monthly(report_data: OrdersMonthlyReportData) -> bytes:
    rl = _require_reportlab()

    colors = rl["colors"]
    TA_CENTER = rl["TA_CENTER"]
    TA_LEFT = rl["TA_LEFT"]
    A4 = rl["A4"]
    landscape = rl["landscape"]
    ParagraphStyle = rl["ParagraphStyle"]
    getSampleStyleSheet = rl["getSampleStyleSheet"]
    mm = rl["mm"]
    PageBreak = rl["PageBreak"]
    Paragraph = rl["Paragraph"]
    SimpleDocTemplate = rl["SimpleDocTemplate"]
    Spacer = rl["Spacer"]
    Table = rl["Table"]
    TableStyle = rl["TableStyle"]

    styles = _build_pdf_styles(ParagraphStyle, getSampleStyleSheet, TA_LEFT, TA_CENTER)

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        rightMargin=10 * mm,
        leftMargin=10 * mm,
        topMargin=16 * mm,
        bottomMargin=12 * mm,
    )

    story = []

    _build_pdf_page_one_orders(
        report_data=report_data,
        story=story,
        Table=Table,
        TableStyle=TableStyle,
        Paragraph=Paragraph,
        Spacer=Spacer,
        colors=colors,
        styles=styles,
    )

    if report_data.sections:
        story.append(PageBreak())

        for section in report_data.sections:
            story.append(_pdf_p(section.section_title, Paragraph, styles["section"]))
            story.append(_build_pdf_orders_table(section, Table, TableStyle, Paragraph, colors, styles))
            story.append(Spacer(1, 8))

        if _should_show_overall_totals(report_data):
            overall = Table(
                [[
                    "OVERALL TOTALS",
                    f"{_format_decimal(report_data.totals.total_quantity_kg)} kg",
                    f"{report_data.totals.total_pieces_required:,} pcs",
                ]],
                colWidths=[130, 140, 120],
            )
            overall.setStyle(
                TableStyle(
                    [
                        ("GRID", (0, 0), (-1, -1), 0.6, colors.black),
                        ("FONTNAME", (0, 0), (-1, -1), "Times-Bold"),
                        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F7F7F7")),
                        ("PADDING", (0, 0), (-1, -1), 5),
                    ]
                )
            )
            story.append(overall)

    doc.build(story, onFirstPage=_pdf_page_decor, onLaterPages=_pdf_page_decor)
    return buffer.getvalue()


def _build_pdf_for_frozen_containers_monthly(
    report_data: FrozenContainersMonthlyReportData,
) -> bytes:
    rl = _require_reportlab()

    colors = rl["colors"]
    TA_CENTER = rl["TA_CENTER"]
    TA_LEFT = rl["TA_LEFT"]
    A4 = rl["A4"]
    landscape = rl["landscape"]
    ParagraphStyle = rl["ParagraphStyle"]
    getSampleStyleSheet = rl["getSampleStyleSheet"]
    mm = rl["mm"]
    Paragraph = rl["Paragraph"]
    SimpleDocTemplate = rl["SimpleDocTemplate"]
    Spacer = rl["Spacer"]
    Table = rl["Table"]
    TableStyle = rl["TableStyle"]

    styles = _build_pdf_styles(ParagraphStyle, getSampleStyleSheet, TA_LEFT, TA_CENTER)

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        rightMargin=8 * mm,
        leftMargin=8 * mm,
        topMargin=16 * mm,
        bottomMargin=12 * mm,
    )

    story = []

    _build_pdf_page_one_frozen(
        report_data=report_data,
        story=story,
        Table=Table,
        TableStyle=TableStyle,
        Paragraph=Paragraph,
        Spacer=Spacer,
        styles=styles,
    )

    story.append(
        _build_pdf_frozen_table(
            report_data.rows,
            report_data.totals,
            Table,
            TableStyle,
            Paragraph,
            colors,
            styles,
        )
    )

    doc.build(story, onFirstPage=_pdf_page_decor, onLaterPages=_pdf_page_decor)
    return buffer.getvalue()


def _build_pdf_for_breakeven_summary(report_data: BreakevenSummaryReportData) -> bytes:
    rl = _require_reportlab()

    colors = rl["colors"]
    TA_CENTER = rl["TA_CENTER"]
    TA_LEFT = rl["TA_LEFT"]
    A4 = rl["A4"]  # portrait
    ParagraphStyle = rl["ParagraphStyle"]
    getSampleStyleSheet = rl["getSampleStyleSheet"]
    mm = rl["mm"]
    Paragraph = rl["Paragraph"]
    SimpleDocTemplate = rl["SimpleDocTemplate"]
    Spacer = rl["Spacer"]
    Table = rl["Table"]
    TableStyle = rl["TableStyle"]

    styles = _build_pdf_styles(ParagraphStyle, getSampleStyleSheet, TA_LEFT, TA_CENTER)

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,  # portrait
        rightMargin=10 * mm,
        leftMargin=10 * mm,
        topMargin=16 * mm,
        bottomMargin=12 * mm,
    )

    story = []

    _build_pdf_page_one_breakeven(
        report_data=report_data,
        story=story,
        Table=Table,
        TableStyle=TableStyle,
        Paragraph=Paragraph,
        Spacer=Spacer,
        styles=styles,
    )

    story.append(
        _build_pdf_breakeven_table(
            _build_breakeven_rows(report_data),
            Table,
            TableStyle,
            Paragraph,
            colors,
            styles,
        )
    )

    doc.build(story, onFirstPage=_pdf_page_decor, onLaterPages=_pdf_page_decor)
    return buffer.getvalue()


# =============================================================================
# DOCX generation
# =============================================================================
def _require_docx():
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError(
            "DOCX generation requires python-docx. Install it with: pip install python-docx"
        ) from exc

    return {
        "Document": Document,
        "WD_ORIENT": WD_ORIENT,
        "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
        "Inches": Inches,
        "Pt": Pt,
        "RGBColor": RGBColor,
    }


def _set_docx_orientation(document, WD_ORIENT, Inches, landscape=True):
    section = document.sections[0]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.left_margin = Inches(0.35)
        section.right_margin = Inches(0.35)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.4)


def _set_docx_base_font(document, Pt):
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)


def _set_run_font(run, name="Times New Roman", size=None, bold=False, color=None):
    run.font.name = name
    if size is not None:
        run.font.size = size
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _left_cell_text(cell):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _set_cell_text_color(cell, RGBColor, rgb_hex="B22222"):
    color = RGBColor.from_string(rgb_hex)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = color
            run.bold = True


def _add_docx_top_bars(document):
    table = document.add_table(rows=1, cols=2)
    table.autofit = True

    left = table.rows[0].cells[0]
    right = table.rows[0].cells[1]
    _shade_cell(left, "2F8F2F")
    _shade_cell(right, "8B0000")

    for cell in table.rows[0].cells:
        p = cell.paragraphs[0]
        run = p.add_run(" ")
        run.font.size = 1


def _add_docx_heading_block(document, report_data, Pt, WD_ALIGN_PARAGRAPH):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(f"{report_data.organization_name or ''} – {report_data.title}")
    _set_run_font(r, size=Pt(18), bold=True)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(
        f"Prepared by: {getattr(report_data, 'prepared_by', '') or ''}    Date: {_format_compact_date(getattr(report_data, 'generated_at', None))}"
    )
    _set_run_font(r, size=Pt(10), bold=True)


def _add_docx_center_heading(document, heading: str, WD_ALIGN_PARAGRAPH, Pt, size=12):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(heading)
    _set_run_font(r, size=Pt(size), bold=True)


def _add_docx_summary_table_from_items(
    document,
    heading: str,
    items: list[tuple[str, str, str]],
    WD_ALIGN_PARAGRAPH,
    Pt,
):
    if not items:
        return

    _add_docx_center_heading(document, heading, WD_ALIGN_PARAGRAPH, Pt, size=12)

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "No"
    hdr[1].text = "Particulars"
    hdr[2].text = "Value"

    for cell in hdr:
        _shade_cell(cell, "EFEFEF")
        _left_cell_text(cell)

    for no, label, value in items:
        row_cells = table.add_row().cells
        row_cells[0].text = no
        row_cells[1].text = label
        row_cells[2].text = value
        _left_cell_text(row_cells[0])
        _left_cell_text(row_cells[1])
        _left_cell_text(row_cells[2])


def _add_docx_projection_table(document, report_data, projection: AnimalProjectionBlock, WD_ALIGN_PARAGRAPH, Pt, RGBColor):
    _add_docx_center_heading(document, _projection_title(report_data), WD_ALIGN_PARAGRAPH, Pt, size=12)

    normalized_rows = _normalize_projection_rows(projection)
    total_goats, total_sheep, total_cattle, grand_total = _projection_totals_from_normalized_rows(normalized_rows)

    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    headers = table.rows[0].cells
    headers[0].text = "WEEK"
    headers[1].text = "GOAT"
    headers[2].text = "SHEEP"
    headers[3].text = "CATTLE"
    headers[4].text = "TOTAL"

    for cell in headers:
        _shade_cell(cell, "EFEFEF")
        _left_cell_text(cell)

    for row in normalized_rows:
        cells = table.add_row().cells
        cells[0].text = str(row["label"])
        cells[1].text = str(row["goats"])
        cells[2].text = str(row["sheep"])
        cells[3].text = str(row["cattle"])
        cells[4].text = str(row["total_animals"])
        for cell in cells:
            _left_cell_text(cell)

    cells = table.add_row().cells
    cells[0].text = "TOTAL (All Animals)"
    cells[1].text = str(total_goats)
    cells[2].text = str(total_sheep)
    cells[3].text = str(total_cattle)
    cells[4].text = str(grand_total)
    for cell in cells:
        _left_cell_text(cell)
        _set_cell_text_color(cell, RGBColor, "B22222")


def _add_docx_orders_section_table(document, section: OrdersReportSection, WD_ALIGN_PARAGRAPH, Pt):
    _add_docx_center_heading(document, section.section_title, WD_ALIGN_PARAGRAPH, Pt, size=11)

    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"

    headers = table.rows[0].cells
    headers[0].text = "Serial No."
    headers[1].text = "Name of Enterprise"
    headers[2].text = "Product Quantity"
    headers[3].text = "Pieces Required"
    headers[4].text = "Slaughter Schedule"
    headers[5].text = "Expected Delivery"
    headers[6].text = "Status"

    for cell in headers:
        _shade_cell(cell, "EFEFEF")
        _left_cell_text(cell)

    for row in section.rows:
        cells = table.add_row().cells
        cells[0].text = f"{row.serial_no}."
        cells[1].text = row.enterprise_name
        cells[2].text = _quantity_text_for_row(row)
        cells[3].text = _pieces_text_for_row(row)
        cells[4].text = _format_human_date(row.slaughter_schedule)
        cells[5].text = _format_human_date(row.expected_delivery)
        cells[6].text = _status_text(row.status)

        for cell in cells:
            _left_cell_text(cell)

    cells = table.add_row().cells
    cells[0].text = "Totals"
    cells[1].text = ""
    cells[2].text = f"{_format_decimal(section.total_quantity_kg)} kg"
    cells[3].text = f"{section.total_pieces_required:,} pcs"
    cells[4].text = ""
    cells[5].text = ""
    cells[6].text = ""

    for cell in cells:
        _left_cell_text(cell)


def _add_docx_frozen_table(document, rows: List[FrozenContainersReportRow], totals):
    table = document.add_table(rows=1, cols=11)
    table.style = "Table Grid"

    headers = table.rows[0].cells
    headers[0].text = "No."
    headers[1].text = "Client Name"
    headers[2].text = "Order Ratio"
    headers[3].text = "Status"
    headers[4].text = "Container Value (USD)"
    headers[5].text = "Price per kg"
    headers[6].text = "Down Payment"
    headers[7].text = "Balance"
    headers[8].text = "Container Gate in"
    headers[9].text = "Departure Date"
    headers[10].text = "Jurisdiction"

    for cell in headers:
        _shade_cell(cell, "EFEFEF")
        _left_cell_text(cell)

    for row in rows:
        cells = table.add_row().cells
        cells[0].text = str(row.serial_no)
        cells[1].text = row.client_name
        cells[2].text = row.order_ratio or ""
        cells[3].text = _status_text(row.status)
        cells[4].text = _format_currency(row.container_value_usd) if row.container_value_usd is not None else "-"
        cells[5].text = _format_decimal4(row.price_per_kg_usd) if row.price_per_kg_usd is not None else "-"
        cells[6].text = _format_currency(row.down_payment_usd) if row.down_payment_usd is not None else "-"
        cells[7].text = _format_currency(row.balance_usd) if row.balance_usd is not None else "-"
        cells[8].text = _format_human_date(row.container_gate_in)
        cells[9].text = _format_human_date(row.departure_date)
        cells[10].text = row.jurisdiction or ""

        for cell in cells:
            _left_cell_text(cell)

    if totals:
        cells = table.add_row().cells
        cells[0].text = "Totals"
        cells[1].text = ""
        cells[2].text = ""
        cells[3].text = ""
        cells[4].text = _format_currency(totals.total_container_value_usd)
        cells[5].text = ""
        cells[6].text = _format_currency(totals.total_down_payment_usd)
        cells[7].text = _format_currency(totals.total_balance_usd)
        cells[8].text = ""
        cells[9].text = ""
        cells[10].text = ""

        for cell in cells:
            _left_cell_text(cell)


def _add_docx_breakeven_table(document, report_data: BreakevenSummaryReportData):
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    headers = table.rows[0].cells
    headers[0].text = "Index"
    headers[1].text = "Metric"
    headers[2].text = "Quantity (Tonnes)"
    headers[3].text = "USD (Total)"
    headers[4].text = "Percentage"

    for cell in headers:
        _shade_cell(cell, "EFEFEF")
        _left_cell_text(cell)

    for row in _build_breakeven_rows(report_data):
        cells = table.add_row().cells
        cells[0].text = str(row.index)
        cells[1].text = row.metric
        cells[2].text = row.quantity_display or ""
        cells[3].text = row.usd_display or ""
        cells[4].text = row.percentage_display or ""
        for cell in cells:
            _left_cell_text(cell)


def _build_docx_for_orders_monthly(report_data: OrdersMonthlyReportData) -> bytes:
    dx = _require_docx()

    Document = dx["Document"]
    WD_ORIENT = dx["WD_ORIENT"]
    WD_ALIGN_PARAGRAPH = dx["WD_ALIGN_PARAGRAPH"]
    Inches = dx["Inches"]
    Pt = dx["Pt"]
    RGBColor = dx["RGBColor"]

    document = Document()
    _set_docx_orientation(document, WD_ORIENT, Inches, landscape=True)
    _set_docx_base_font(document, Pt)

    _add_docx_top_bars(document)
    _add_docx_heading_block(document, report_data, Pt, WD_ALIGN_PARAGRAPH)
    _add_docx_summary_table_from_items(
        document,
        _summary_title(report_data),
        _build_orders_summary_items(report_data),
        WD_ALIGN_PARAGRAPH,
        Pt,
    )

    if report_data.animal_projection is not None:
        _add_docx_projection_table(
            document,
            report_data,
            report_data.animal_projection,
            WD_ALIGN_PARAGRAPH,
            Pt,
            RGBColor,
        )

    if report_data.sections:
        document.add_page_break()
        for section in report_data.sections:
            _add_docx_orders_section_table(document, section, WD_ALIGN_PARAGRAPH, Pt)

        if _should_show_overall_totals(report_data):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(
                f"OVERALL TOTALS: {_format_decimal(report_data.totals.total_quantity_kg)} kg    {report_data.totals.total_pieces_required:,} pcs"
            )
            _set_run_font(r, bold=True)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_docx_for_frozen_containers_monthly(
    report_data: FrozenContainersMonthlyReportData,
) -> bytes:
    dx = _require_docx()

    Document = dx["Document"]
    WD_ORIENT = dx["WD_ORIENT"]
    WD_ALIGN_PARAGRAPH = dx["WD_ALIGN_PARAGRAPH"]
    Inches = dx["Inches"]
    Pt = dx["Pt"]

    document = Document()
    _set_docx_orientation(document, WD_ORIENT, Inches, landscape=True)
    _set_docx_base_font(document, Pt)

    _add_docx_top_bars(document)
    _add_docx_heading_block(document, report_data, Pt, WD_ALIGN_PARAGRAPH)
    _add_docx_frozen_table(document, report_data.rows, report_data.totals)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_docx_for_breakeven_summary(report_data: BreakevenSummaryReportData) -> bytes:
    dx = _require_docx()

    Document = dx["Document"]
    WD_ORIENT = dx["WD_ORIENT"]
    WD_ALIGN_PARAGRAPH = dx["WD_ALIGN_PARAGRAPH"]
    Inches = dx["Inches"]
    Pt = dx["Pt"]

    document = Document()
    _set_docx_orientation(document, WD_ORIENT, Inches, landscape=False)  # portrait
    _set_docx_base_font(document, Pt)

    _add_docx_top_bars(document)
    _add_docx_heading_block(document, report_data, Pt, WD_ALIGN_PARAGRAPH)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(f"Report Date: {_format_report_day(report_data.report_date)}")
    _set_run_font(r, size=Pt(10), bold=True)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(f"Month/Year: {report_data.month or ''}/{report_data.year or ''}")
    _set_run_font(r, size=Pt(10), bold=True)

    _add_docx_breakeven_table(document, report_data)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# =============================================================================
# Public API
# =============================================================================
def generate_report(
    report_data: ReportData,
    output_format: str,
) -> GeneratedReport:
    output_format = (output_format or "").strip().lower()

    if report_data.report_type == "orders_monthly":
        if output_format == "csv":
            content = _build_csv_for_orders_monthly(report_data)
        elif output_format == "pdf":
            content = _build_pdf_for_orders_monthly(report_data)
        elif output_format == "docx":
            content = _build_docx_for_orders_monthly(report_data)
        else:
            raise ValueError("Unsupported report format. Use csv, pdf, or docx.")

    elif report_data.report_type == "frozen_containers_monthly":
        if output_format == "csv":
            content = _build_csv_for_frozen_containers_monthly(report_data)
        elif output_format == "pdf":
            content = _build_pdf_for_frozen_containers_monthly(report_data)
        elif output_format == "docx":
            content = _build_docx_for_frozen_containers_monthly(report_data)
        else:
            raise ValueError("Unsupported report format. Use csv, pdf, or docx.")

    elif report_data.report_type == "breakeven_summary":
        if output_format == "csv":
            content = _build_csv_for_breakeven_summary(report_data)
        elif output_format == "pdf":
            content = _build_pdf_for_breakeven_summary(report_data)
        elif output_format == "docx":
            content = _build_docx_for_breakeven_summary(report_data)
        else:
            raise ValueError("Unsupported report format. Use csv, pdf, or docx.")

    else:
        raise ValueError(f"Unsupported report type: {report_data.report_type}")

    media_type = {
        "csv": CSV_MEDIA_TYPE,
        "pdf": PDF_MEDIA_TYPE,
        "docx": DOCX_MEDIA_TYPE,
    }[output_format]

    filename = _build_report_filename(
        report_type=report_data.report_type,
        report_format=output_format,
    )

    return GeneratedReport(
        filename=filename,
        media_type=media_type,
        content=content,
    )