from __future__ import annotations

import csv
import math
from collections import defaultdict
from dataclasses import dataclass
from datetime import date, datetime
from decimal import Decimal
from io import BytesIO, StringIO
from typing import Any

from sqlalchemy import func, or_
from sqlalchemy.exc import IntegrityError
from sqlalchemy.orm import Session, joinedload

from app.models.inventory import (
    ConsumableStoreInventory,
    ConsumableStoreName,
    InventoryConsumableCategory,
    InventoryConsumableItem,
    InventoryProduct,
    InventoryProductCategory,
    ProductBalanceUnit,
    ProductStoreInventory,
    ProductStoreName,
)
from app.schemas.inventory import (
    ConsumableInventoryReportGroup,
    ConsumableInventoryReportRequest,
    ConsumableInventoryReportResponse,
    ConsumableInventoryReportRow,
    ConsumableInventoryReportSummary,
    ConsumableStoreInventoryCreate,
    ConsumableStoreInventoryFilter,
    ConsumableStoreInventoryListResponse,
    ConsumableStoreInventoryRead,
    ConsumableStoreInventoryUpdate,
    InventoryBootstrapResponse,
    InventoryConsumableCategoryCreate,
    InventoryConsumableCategoryRead,
    InventoryConsumableCategoryUpdate,
    InventoryConsumableItemCreate,
    InventoryConsumableItemRead,
    InventoryConsumableItemUpdate,
    InventoryExportFileResponse,
    InventoryProductCategoryCreate,
    InventoryProductCategoryRead,
    InventoryProductCategoryUpdate,
    InventoryProductCreate,
    InventoryProductRead,
    InventoryProductUpdate,
    InventoryStoreOptionSchema,
    ProductInventoryReportGroup,
    ProductInventoryReportRequest,
    ProductInventoryReportResponse,
    ProductInventoryReportRow,
    ProductInventoryReportSummary,
    ProductStoreInventoryCreate,
    ProductStoreInventoryFilter,
    ProductStoreInventoryListResponse,
    ProductStoreInventoryRead,
    ProductStoreInventoryUpdate,
)

TWOPLACES = Decimal("0.01")


@dataclass
class GeneratedReportFile:
    filename: str
    content_type: str
    data: bytes

    def to_schema(self, report_type: str, export_format: str) -> InventoryExportFileResponse:
        return InventoryExportFileResponse(
            filename=self.filename,
            content_type=self.content_type,
            report_type=report_type,
            export_format=export_format,
            generated_at=datetime.utcnow(),
        )


class InventoryServiceError(Exception):
    pass


class InventoryNotFoundError(InventoryServiceError):
    pass


class InventoryConflictError(InventoryServiceError):
    pass


def qty(value: Any) -> Decimal:
    if value is None:
        return Decimal("0.00")
    if isinstance(value, Decimal):
        return value.quantize(TWOPLACES)
    return Decimal(str(value)).quantize(TWOPLACES)


def fmt_qty(value: Any) -> str:
    return f"{qty(value):,.2f}"


def fmt_date(value: date | None) -> str:
    return value.strftime("%d-%b-%Y") if value else ""


def safe_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value)


def decimal_equal(a: Any, b: Any) -> bool:
    return qty(a) == qty(b)


class InventoryService:
    def __init__(self, db: Session):
        self.db = db

    # =========================================================================
    # BOOTSTRAP / DROPDOWNS
    # =========================================================================

    def get_bootstrap_data(self) -> InventoryBootstrapResponse:
        return InventoryBootstrapResponse(
            product_stores=[
                InventoryStoreOptionSchema(value=store.value, label=store.value)
                for store in ProductStoreName
            ],
            consumable_stores=[
                InventoryStoreOptionSchema(value=store.value, label=store.value)
                for store in ConsumableStoreName
            ],
            product_categories=self.list_product_categories(),
            products=self.list_products(),
            consumable_categories=self.list_consumable_categories(),
            consumable_items=self.list_consumable_items(),
        )

    # =========================================================================
    # PRODUCT CATEGORY CRUD
    # =========================================================================

    def list_product_categories(self, active_only: bool = False) -> list[InventoryProductCategoryRead]:
        query = self.db.query(InventoryProductCategory).order_by(
            InventoryProductCategory.sort_order.asc(),
            InventoryProductCategory.name.asc(),
        )
        if active_only:
            query = query.filter(InventoryProductCategory.is_active.is_(True))
        return [InventoryProductCategoryRead.model_validate(row) for row in query.all()]

    def get_product_category_or_404(self, category_id: str) -> InventoryProductCategory:
        row = (
            self.db.query(InventoryProductCategory)
            .filter(InventoryProductCategory.id == category_id)
            .first()
        )
        if not row:
            raise InventoryNotFoundError("Product category not found.")
        return row

    def create_product_category(self, payload: InventoryProductCategoryCreate) -> InventoryProductCategoryRead:
        row = InventoryProductCategory(**payload.model_dump())
        self.db.add(row)
        self._commit_or_conflict("A product category with this name already exists.")
        self.db.refresh(row)
        return InventoryProductCategoryRead.model_validate(row)

    def update_product_category(
        self,
        category_id: str,
        payload: InventoryProductCategoryUpdate,
    ) -> InventoryProductCategoryRead:
        row = self.get_product_category_or_404(category_id)
        for key, value in payload.model_dump(exclude_unset=True).items():
            setattr(row, key, value)
        self._commit_or_conflict("A product category with this name already exists.")
        self.db.refresh(row)
        return InventoryProductCategoryRead.model_validate(row)

    def delete_product_category(self, category_id: str) -> None:
        row = self.get_product_category_or_404(category_id)
        self.db.delete(row)
        self._commit_or_conflict("Unable to delete product category because it is in use.")

    # =========================================================================
    # PRODUCT CRUD
    # =========================================================================

    def list_products(self, active_only: bool = False) -> list[InventoryProductRead]:
        query = (
            self.db.query(InventoryProduct)
            .options(joinedload(InventoryProduct.category))
            .order_by(InventoryProduct.sort_order.asc(), InventoryProduct.name.asc())
        )
        if active_only:
            query = query.filter(InventoryProduct.is_active.is_(True))

        rows = query.all()
        return [
            InventoryProductRead.model_validate(
                {
                    **row.__dict__,
                    "category_name": row.category.name if row.category else None,
                }
            )
            for row in rows
        ]

    def get_product_or_404(self, product_id: str) -> InventoryProduct:
        row = (
            self.db.query(InventoryProduct)
            .options(joinedload(InventoryProduct.category))
            .filter(InventoryProduct.id == product_id)
            .first()
        )
        if not row:
            raise InventoryNotFoundError("Product not found.")
        return row

    def create_product(self, payload: InventoryProductCreate) -> InventoryProductRead:
        self.get_product_category_or_404(payload.category_id)
        row = InventoryProduct(**payload.model_dump())
        self.db.add(row)
        self._commit_or_conflict("A product with this name already exists in that category.")
        self.db.refresh(row)
        row = self.get_product_or_404(row.id)
        return InventoryProductRead.model_validate(
            {
                **row.__dict__,
                "category_name": row.category.name if row.category else None,
            }
        )

    def update_product(self, product_id: str, payload: InventoryProductUpdate) -> InventoryProductRead:
        row = self.get_product_or_404(product_id)
        data = payload.model_dump(exclude_unset=True)

        if "category_id" in data and data["category_id"]:
            self.get_product_category_or_404(data["category_id"])

        for key, value in data.items():
            setattr(row, key, value)

        self._commit_or_conflict("A product with this name already exists in that category.")
        self.db.refresh(row)
        row = self.get_product_or_404(row.id)

        return InventoryProductRead.model_validate(
            {
                **row.__dict__,
                "category_name": row.category.name if row.category else None,
            }
        )

    def delete_product(self, product_id: str) -> None:
        row = self.get_product_or_404(product_id)
        self.db.delete(row)
        self._commit_or_conflict("Unable to delete product because it is in use.")

    # =========================================================================
    # CONSUMABLE CATEGORY CRUD
    # =========================================================================

    def list_consumable_categories(self, active_only: bool = False) -> list[InventoryConsumableCategoryRead]:
        query = self.db.query(InventoryConsumableCategory).order_by(
            InventoryConsumableCategory.sort_order.asc(),
            InventoryConsumableCategory.name.asc(),
        )
        if active_only:
            query = query.filter(InventoryConsumableCategory.is_active.is_(True))
        return [InventoryConsumableCategoryRead.model_validate(row) for row in query.all()]

    def get_consumable_category_or_404(self, category_id: str) -> InventoryConsumableCategory:
        row = (
            self.db.query(InventoryConsumableCategory)
            .filter(InventoryConsumableCategory.id == category_id)
            .first()
        )
        if not row:
            raise InventoryNotFoundError("Consumable category not found.")
        return row

    def create_consumable_category(
        self,
        payload: InventoryConsumableCategoryCreate,
    ) -> InventoryConsumableCategoryRead:
        row = InventoryConsumableCategory(**payload.model_dump())
        self.db.add(row)
        self._commit_or_conflict("A consumable category with this name already exists.")
        self.db.refresh(row)
        return InventoryConsumableCategoryRead.model_validate(row)

    def update_consumable_category(
        self,
        category_id: str,
        payload: InventoryConsumableCategoryUpdate,
    ) -> InventoryConsumableCategoryRead:
        row = self.get_consumable_category_or_404(category_id)
        for key, value in payload.model_dump(exclude_unset=True).items():
            setattr(row, key, value)
        self._commit_or_conflict("A consumable category with this name already exists.")
        self.db.refresh(row)
        return InventoryConsumableCategoryRead.model_validate(row)

    def delete_consumable_category(self, category_id: str) -> None:
        row = self.get_consumable_category_or_404(category_id)
        self.db.delete(row)
        self._commit_or_conflict("Unable to delete consumable category because it is in use.")

    # =========================================================================
    # CONSUMABLE ITEM CRUD
    # =========================================================================

    def list_consumable_items(self, active_only: bool = False) -> list[InventoryConsumableItemRead]:
        query = (
            self.db.query(InventoryConsumableItem)
            .options(joinedload(InventoryConsumableItem.category))
            .order_by(InventoryConsumableItem.sort_order.asc(), InventoryConsumableItem.name.asc())
        )
        if active_only:
            query = query.filter(InventoryConsumableItem.is_active.is_(True))

        rows = query.all()
        return [
            InventoryConsumableItemRead.model_validate(
                {
                    **row.__dict__,
                    "category_name": row.category.name if row.category else None,
                }
            )
            for row in rows
        ]

    def get_consumable_item_or_404(self, item_id: str) -> InventoryConsumableItem:
        row = (
            self.db.query(InventoryConsumableItem)
            .options(joinedload(InventoryConsumableItem.category))
            .filter(InventoryConsumableItem.id == item_id)
            .first()
        )
        if not row:
            raise InventoryNotFoundError("Consumable item not found.")
        return row

    def create_consumable_item(self, payload: InventoryConsumableItemCreate) -> InventoryConsumableItemRead:
        self.get_consumable_category_or_404(payload.category_id)
        row = InventoryConsumableItem(**payload.model_dump())
        self.db.add(row)
        self._commit_or_conflict("A consumable item with this name already exists in that category.")
        self.db.refresh(row)
        row = self.get_consumable_item_or_404(row.id)

        return InventoryConsumableItemRead.model_validate(
            {
                **row.__dict__,
                "category_name": row.category.name if row.category else None,
            }
        )

    def update_consumable_item(
        self,
        item_id: str,
        payload: InventoryConsumableItemUpdate,
    ) -> InventoryConsumableItemRead:
        row = self.get_consumable_item_or_404(item_id)
        data = payload.model_dump(exclude_unset=True)

        if "category_id" in data and data["category_id"]:
            self.get_consumable_category_or_404(data["category_id"])

        for key, value in data.items():
            setattr(row, key, value)

        self._commit_or_conflict("A consumable item with this name already exists in that category.")
        self.db.refresh(row)
        row = self.get_consumable_item_or_404(row.id)

        return InventoryConsumableItemRead.model_validate(
            {
                **row.__dict__,
                "category_name": row.category.name if row.category else None,
            }
        )

    def delete_consumable_item(self, item_id: str) -> None:
        row = self.get_consumable_item_or_404(item_id)
        self.db.delete(row)
        self._commit_or_conflict("Unable to delete consumable item because it is in use.")

    # =========================================================================
    # OPENING BALANCE PREVIEW HELPERS
    # =========================================================================

    def get_product_opening_balance_preview(
        self,
        *,
        entry_date: date,
        store: ProductStoreName | str,
        product_id: str,
    ) -> Decimal:
        store_value = self._normalize_product_store(store)
        return ProductStoreInventory.get_opening_balance_for_date(
            self.db,
            entry_date=entry_date,
            store=store_value,
            product_id=product_id,
        )

    def get_consumable_opening_balance_preview(
        self,
        *,
        entry_date: date,
        store: ConsumableStoreName | str,
        item_id: str,
    ) -> Decimal:
        store_value = self._normalize_consumable_store(store)
        return ConsumableStoreInventory.get_opening_balance_for_date(
            self.db,
            entry_date=entry_date,
            store=store_value,
            item_id=item_id,
        )

    # =========================================================================
    # DAILY SHEET BUILDERS (FIXED)
    # =========================================================================

    def get_product_daily_sheet(
        self,
        *,
        entry_date: date,
        store: ProductStoreName | str,
        active_only: bool = True,
    ) -> dict[str, Any]:
        store_value = self._normalize_product_store(store)
        products = self._list_product_models(active_only=active_only)
        existing_rows = self._list_product_entries_for_sheet(entry_date=entry_date, store=store_value)
        existing_by_product_id = {row.product_id: row for row in existing_rows}

        rows: list[dict[str, Any]] = []
        serial_cursor = 1
        saved_count = 0
        generated_count = 0

        for product in products:
            saved_row = existing_by_product_id.pop(product.id, None)
            if saved_row:
                rows.append(self._product_sheet_row_from_existing(saved_row))
                saved_count += 1
                serial_cursor = max(serial_cursor, int(saved_row.serial_no or 0) + 1)
            else:
                rows.append(
                    self._product_sheet_row_from_product(
                        product=product,
                        entry_date=entry_date,
                        store=store_value,
                        serial_no=serial_cursor,
                    )
                )
                generated_count += 1
                serial_cursor += 1

        leftover_saved_rows = sorted(
            existing_by_product_id.values(),
            key=lambda row: (
                int(row.serial_no or 0),
                safe_text(row.product_category_name).lower(),
                safe_text(row.product_name).lower(),
            ),
        )
        for saved_row in leftover_saved_rows:
            rows.append(self._product_sheet_row_from_existing(saved_row))
            saved_count += 1

        # FIX: Reassign serial numbers sequentially based on final order
        for idx, row in enumerate(rows, start=1):
            row["serial_no"] = idx

        return {
            "entry_date": entry_date,
            "store": store_value,
            "total_rows": len(rows),
            "saved_rows": saved_count,
            "generated_rows": generated_count,
            "has_saved_rows": saved_count > 0,
            "rows": rows,
        }

    def get_consumable_daily_sheet(
        self,
        *,
        entry_date: date,
        store: ConsumableStoreName | str,
        active_only: bool = True,
    ) -> dict[str, Any]:
        store_value = self._normalize_consumable_store(store)
        items = self._list_consumable_item_models(active_only=active_only)
        existing_rows = self._list_consumable_entries_for_sheet(entry_date=entry_date, store=store_value)
        existing_by_item_id = {row.item_id: row for row in existing_rows}

        rows: list[dict[str, Any]] = []
        serial_cursor = 1
        saved_count = 0
        generated_count = 0

        for item in items:
            saved_row = existing_by_item_id.pop(item.id, None)
            if saved_row:
                rows.append(self._consumable_sheet_row_from_existing(saved_row))
                saved_count += 1
                serial_cursor = max(serial_cursor, int(saved_row.serial_no or 0) + 1)
            else:
                rows.append(
                    self._consumable_sheet_row_from_item(
                        item=item,
                        entry_date=entry_date,
                        store=store_value,
                        serial_no=serial_cursor,
                    )
                )
                generated_count += 1
                serial_cursor += 1

        leftover_saved_rows = sorted(
            existing_by_item_id.values(),
            key=lambda row: (
                int(row.serial_no or 0),
                safe_text(row.item_category_name).lower(),
                safe_text(row.item_name).lower(),
            ),
        )
        for saved_row in leftover_saved_rows:
            rows.append(self._consumable_sheet_row_from_existing(saved_row))
            saved_count += 1

        # FIX: Reassign serial numbers sequentially based on final order
        for idx, row in enumerate(rows, start=1):
            row["serial_no"] = idx

        return {
            "entry_date": entry_date,
            "store": store_value,
            "total_rows": len(rows),
            "saved_rows": saved_count,
            "generated_rows": generated_count,
            "has_saved_rows": saved_count > 0,
            "rows": rows,
        }

    # =========================================================================
    # PRODUCT INVENTORY CRUD
    # =========================================================================

    def get_product_inventory_or_404(self, entry_id: str) -> ProductStoreInventory:
        row = (
            self.db.query(ProductStoreInventory)
            .options(
                joinedload(ProductStoreInventory.product_category),
                joinedload(ProductStoreInventory.product),
            )
            .filter(ProductStoreInventory.id == entry_id)
            .first()
        )
        if not row:
            raise InventoryNotFoundError("Product inventory entry not found.")
        return row

    def create_product_inventory_entry(
        self,
        payload: ProductStoreInventoryCreate,
        *,
        created_by: str | None = None,
    ) -> ProductStoreInventoryRead:
        row = self._build_new_product_inventory_row(payload)
        row.created_by = created_by
        row.updated_by = created_by

        if row.serial_no is None:
            row.serial_no = self._next_product_serial_no(row.entry_date, row.store)

        self._prepare_product_inventory_row(
            row,
            overwrite_opening_balance=payload.overwrite_opening_balance,
        )

        self.db.add(row)
        self._commit_or_conflict(
            "A product inventory entry already exists for this date, store, and product."
        )
        self.db.refresh(row)
        return self._to_product_inventory_read(row)

    def update_product_inventory_entry(
        self,
        entry_id: str,
        payload: ProductStoreInventoryUpdate,
        *,
        updated_by: str | None = None,
    ) -> ProductStoreInventoryRead:
        row = self.get_product_inventory_or_404(entry_id)
        data = payload.model_dump(exclude_unset=True, exclude={"overwrite_opening_balance"})

        for key, value in data.items():
            setattr(row, key, value)

        if row.serial_no is None:
            row.serial_no = self._next_product_serial_no(row.entry_date, row.store)

        row.updated_by = updated_by or row.updated_by

        self._prepare_product_inventory_row(
            row,
            overwrite_opening_balance=payload.overwrite_opening_balance,
            exclude_entry_id=row.id,
        )

        self._commit_or_conflict(
            "A product inventory entry already exists for this date, store, and product."
        )
        self.db.refresh(row)
        return self._to_product_inventory_read(row)

    def delete_product_inventory_entry(self, entry_id: str) -> None:
        row = self.get_product_inventory_or_404(entry_id)
        self.db.delete(row)
        self._commit_or_conflict("Unable to delete product inventory entry.")

    def list_product_inventory(
        self,
        filters: ProductStoreInventoryFilter,
    ) -> ProductStoreInventoryListResponse:
        query = self.db.query(ProductStoreInventory)

        if filters.start_date:
            query = query.filter(ProductStoreInventory.entry_date >= filters.start_date)
        if filters.end_date:
            query = query.filter(ProductStoreInventory.entry_date <= filters.end_date)
        if filters.store:
            query = query.filter(ProductStoreInventory.store == filters.store)
        if filters.product_category_id:
            query = query.filter(ProductStoreInventory.product_category_id == filters.product_category_id)
        if filters.product_id:
            query = query.filter(ProductStoreInventory.product_id == filters.product_id)
        if filters.balance_unit:
            query = query.filter(ProductStoreInventory.balance_unit == filters.balance_unit)
        if filters.checked_by_initials:
            query = query.filter(ProductStoreInventory.checked_by_initials == filters.checked_by_initials)
        if filters.search:
            term = f"%{filters.search}%"
            query = query.filter(
                or_(
                    ProductStoreInventory.product_category_name.ilike(term),
                    ProductStoreInventory.product_name.ilike(term),
                    ProductStoreInventory.remarks.ilike(term),
                    ProductStoreInventory.checked_by_initials.ilike(term),
                )
            )

        total = query.count()
        items = (
            query.order_by(
                ProductStoreInventory.entry_date.desc(),
                ProductStoreInventory.store.asc(),
                ProductStoreInventory.serial_no.asc(),
                ProductStoreInventory.created_at.asc(),
            )
            .offset((filters.page - 1) * filters.page_size)
            .limit(filters.page_size)
            .all()
        )

        return ProductStoreInventoryListResponse(
            total=total,
            page=filters.page,
            page_size=filters.page_size,
            total_pages=math.ceil(total / filters.page_size) if total else 0,
            items=[self._to_product_inventory_read(row) for row in items],
        )

    # =========================================================================
    # PRODUCT BULK UPSERT
    # =========================================================================

    def bulk_upsert_product_inventory(
        self,
        *,
        entry_date: date,
        store: ProductStoreName | str,
        rows: list[dict[str, Any] | ProductStoreInventoryCreate],
        actor_id: str | None = None,
    ) -> dict[str, Any]:
        store_value = self._normalize_product_store(store)
        self._validate_no_duplicate_product_rows(rows)

        saved_rows: list[ProductStoreInventory] = []
        created_count = 0
        updated_count = 0

        try:
            serial_cursor = self._next_product_serial_no(entry_date, store_value)

            for row_index, raw_row in enumerate(rows, start=1):
                payload, provided_fields = self._coerce_product_bulk_payload(
                    raw_row=raw_row,
                    entry_date=entry_date,
                    store=store_value,
                    default_serial_no=serial_cursor,
                )

                product = self.get_product_or_404(payload.product_id)
                existing = self._get_existing_product_entry(
                    entry_date=entry_date,
                    store=store_value,
                    product_id=payload.product_id,
                )

                try:
                    if existing:
                        self._apply_product_bulk_payload_to_existing_row(
                            row=existing,
                            payload=payload,
                            provided_fields=provided_fields,
                            actor_id=actor_id,
                            fallback_serial_no=serial_cursor,
                        )
                        self._prepare_product_inventory_row(
                            existing,
                            overwrite_opening_balance=payload.overwrite_opening_balance,
                            exclude_entry_id=existing.id,
                        )
                        self.db.flush()
                        saved_rows.append(existing)
                        updated_count += 1
                    else:
                        new_row = self._build_new_product_inventory_row(payload)
                        if new_row.serial_no is None:
                            new_row.serial_no = serial_cursor
                        new_row.created_by = actor_id
                        new_row.updated_by = actor_id

                        self._prepare_product_inventory_row(
                            new_row,
                            overwrite_opening_balance=payload.overwrite_opening_balance,
                        )
                        self.db.add(new_row)
                        self.db.flush()
                        saved_rows.append(new_row)
                        created_count += 1

                    serial_cursor += 1
                except InventoryServiceError as exc:
                    product_name = product.name if product else payload.product_id
                    raise InventoryServiceError(
                        f"Product row {row_index} ({product_name}): {exc}"
                    ) from exc

            self.db.commit()

            for row in saved_rows:
                self.db.refresh(row)

            return {
                "entry_date": entry_date,
                "store": store_value,
                "total_rows": len(saved_rows),
                "created_count": created_count,
                "updated_count": updated_count,
                "items": [self._to_product_inventory_read(row).model_dump(mode="python") for row in saved_rows],
            }
        except IntegrityError as exc:
            self.db.rollback()
            raise InventoryConflictError(
                "One or more product inventory rows conflict with existing records."
            ) from exc
        except Exception:
            self.db.rollback()
            raise

    # =========================================================================
    # CONSUMABLE INVENTORY CRUD
    # =========================================================================

    def get_consumable_inventory_or_404(self, entry_id: str) -> ConsumableStoreInventory:
        row = (
            self.db.query(ConsumableStoreInventory)
            .options(
                joinedload(ConsumableStoreInventory.item_category),
                joinedload(ConsumableStoreInventory.item),
            )
            .filter(ConsumableStoreInventory.id == entry_id)
            .first()
        )
        if not row:
            raise InventoryNotFoundError("Consumable inventory entry not found.")
        return row

    def create_consumable_inventory_entry(
        self,
        payload: ConsumableStoreInventoryCreate,
        *,
        created_by: str | None = None,
    ) -> ConsumableStoreInventoryRead:
        row = self._build_new_consumable_inventory_row(payload)
        row.created_by = created_by
        row.updated_by = created_by

        if row.serial_no is None:
            row.serial_no = self._next_consumable_serial_no(row.entry_date, row.store)

        self._prepare_consumable_inventory_row(
            row,
            overwrite_opening_balance=payload.overwrite_opening_balance,
        )

        self.db.add(row)
        self._commit_or_conflict(
            "A consumable inventory entry already exists for this date, store, and item."
        )
        self.db.refresh(row)
        return self._to_consumable_inventory_read(row)

    def update_consumable_inventory_entry(
        self,
        entry_id: str,
        payload: ConsumableStoreInventoryUpdate,
        *,
        updated_by: str | None = None,
    ) -> ConsumableStoreInventoryRead:
        row = self.get_consumable_inventory_or_404(entry_id)
        data = payload.model_dump(exclude_unset=True, exclude={"overwrite_opening_balance"})

        for key, value in data.items():
            setattr(row, key, value)

        if row.serial_no is None:
            row.serial_no = self._next_consumable_serial_no(row.entry_date, row.store)

        row.updated_by = updated_by or row.updated_by

        self._prepare_consumable_inventory_row(
            row,
            overwrite_opening_balance=payload.overwrite_opening_balance,
            exclude_entry_id=row.id,
        )

        self._commit_or_conflict(
            "A consumable inventory entry already exists for this date, store, and item."
        )
        self.db.refresh(row)
        return self._to_consumable_inventory_read(row)

    def delete_consumable_inventory_entry(self, entry_id: str) -> None:
        row = self.get_consumable_inventory_or_404(entry_id)
        self.db.delete(row)
        self._commit_or_conflict("Unable to delete consumable inventory entry.")

    def list_consumable_inventory(
        self,
        filters: ConsumableStoreInventoryFilter,
    ) -> ConsumableStoreInventoryListResponse:
        query = self.db.query(ConsumableStoreInventory)

        if filters.start_date:
            query = query.filter(ConsumableStoreInventory.entry_date >= filters.start_date)
        if filters.end_date:
            query = query.filter(ConsumableStoreInventory.entry_date <= filters.end_date)
        if filters.store:
            query = query.filter(ConsumableStoreInventory.store == filters.store)
        if filters.item_category_id:
            query = query.filter(ConsumableStoreInventory.item_category_id == filters.item_category_id)
        if filters.item_id:
            query = query.filter(ConsumableStoreInventory.item_id == filters.item_id)
        if filters.unit:
            query = query.filter(ConsumableStoreInventory.unit == filters.unit)
        if filters.checked_by_initials:
            query = query.filter(ConsumableStoreInventory.checked_by_initials == filters.checked_by_initials)
        if filters.search:
            term = f"%{filters.search}%"
            query = query.filter(
                or_(
                    ConsumableStoreInventory.item_category_name.ilike(term),
                    ConsumableStoreInventory.item_name.ilike(term),
                    ConsumableStoreInventory.unit.ilike(term),
                    ConsumableStoreInventory.remarks.ilike(term),
                    ConsumableStoreInventory.checked_by_initials.ilike(term),
                )
            )

        total = query.count()
        items = (
            query.order_by(
                ConsumableStoreInventory.entry_date.desc(),
                ConsumableStoreInventory.store.asc(),
                ConsumableStoreInventory.serial_no.asc(),
                ConsumableStoreInventory.created_at.asc(),
            )
            .offset((filters.page - 1) * filters.page_size)
            .limit(filters.page_size)
            .all()
        )

        return ConsumableStoreInventoryListResponse(
            total=total,
            page=filters.page,
            page_size=filters.page_size,
            total_pages=math.ceil(total / filters.page_size) if total else 0,
            items=[self._to_consumable_inventory_read(row) for row in items],
        )

    # =========================================================================
    # CONSUMABLE BULK UPSERT
    # =========================================================================

    def bulk_upsert_consumable_inventory(
        self,
        *,
        entry_date: date,
        store: ConsumableStoreName | str,
        rows: list[dict[str, Any] | ConsumableStoreInventoryCreate],
        actor_id: str | None = None,
    ) -> dict[str, Any]:
        store_value = self._normalize_consumable_store(store)
        self._validate_no_duplicate_consumable_rows(rows)

        saved_rows: list[ConsumableStoreInventory] = []
        created_count = 0
        updated_count = 0

        try:
            serial_cursor = self._next_consumable_serial_no(entry_date, store_value)

            for row_index, raw_row in enumerate(rows, start=1):
                payload, provided_fields = self._coerce_consumable_bulk_payload(
                    raw_row=raw_row,
                    entry_date=entry_date,
                    store=store_value,
                    default_serial_no=serial_cursor,
                )

                item = self.get_consumable_item_or_404(payload.item_id)
                existing = self._get_existing_consumable_entry(
                    entry_date=entry_date,
                    store=store_value,
                    item_id=payload.item_id,
                )

                try:
                    if existing:
                        self._apply_consumable_bulk_payload_to_existing_row(
                            row=existing,
                            payload=payload,
                            provided_fields=provided_fields,
                            actor_id=actor_id,
                            fallback_serial_no=serial_cursor,
                        )
                        self._prepare_consumable_inventory_row(
                            existing,
                            overwrite_opening_balance=payload.overwrite_opening_balance,
                            exclude_entry_id=existing.id,
                        )
                        self.db.flush()
                        saved_rows.append(existing)
                        updated_count += 1
                    else:
                        new_row = self._build_new_consumable_inventory_row(payload)
                        if new_row.serial_no is None:
                            new_row.serial_no = serial_cursor
                        new_row.created_by = actor_id
                        new_row.updated_by = actor_id

                        self._prepare_consumable_inventory_row(
                            new_row,
                            overwrite_opening_balance=payload.overwrite_opening_balance,
                        )
                        self.db.add(new_row)
                        self.db.flush()
                        saved_rows.append(new_row)
                        created_count += 1

                    serial_cursor += 1
                except InventoryServiceError as exc:
                    item_name = item.name if item else payload.item_id
                    raise InventoryServiceError(
                        f"Consumable row {row_index} ({item_name}): {exc}"
                    ) from exc

            self.db.commit()

            for row in saved_rows:
                self.db.refresh(row)

            return {
                "entry_date": entry_date,
                "store": store_value,
                "total_rows": len(saved_rows),
                "created_count": created_count,
                "updated_count": updated_count,
                "items": [self._to_consumable_inventory_read(row).model_dump(mode="python") for row in saved_rows],
            }
        except IntegrityError as exc:
            self.db.rollback()
            raise InventoryConflictError(
                "One or more consumable inventory rows conflict with existing records."
            ) from exc
        except Exception:
            self.db.rollback()
            raise

    # =========================================================================
    # PRODUCT REPORTS / EXPORT
    # =========================================================================

    def build_product_inventory_report(
        self,
        payload: ProductInventoryReportRequest,
    ) -> ProductInventoryReportResponse:
        query = self.db.query(ProductStoreInventory).filter(
            ProductStoreInventory.entry_date >= payload.start_date,
            ProductStoreInventory.entry_date <= payload.end_date,
        )

        if payload.store:
            query = query.filter(ProductStoreInventory.store == payload.store)
        if payload.product_category_id:
            query = query.filter(ProductStoreInventory.product_category_id == payload.product_category_id)
        if payload.product_id:
            query = query.filter(ProductStoreInventory.product_id == payload.product_id)
        if payload.balance_unit:
            query = query.filter(ProductStoreInventory.balance_unit == payload.balance_unit)
        if payload.checked_by_initials:
            query = query.filter(ProductStoreInventory.checked_by_initials == payload.checked_by_initials)

        rows = query.order_by(
            ProductStoreInventory.entry_date.asc(),
            ProductStoreInventory.store.asc(),
            ProductStoreInventory.product_category_name.asc(),
            ProductStoreInventory.product_name.asc(),
            ProductStoreInventory.serial_no.asc(),
            ProductStoreInventory.created_at.asc(),
        ).all()

        groups = self._group_product_report_rows(rows, payload.report_type)
        grand_totals = ProductInventoryReportSummary()

        for group in groups:
            grand_totals.opening_balance += qty(group.summary.opening_balance)
            grand_totals.inflow_production += qty(group.summary.inflow_production)
            grand_totals.inflow_transfers_in += qty(group.summary.inflow_transfers_in)
            grand_totals.outflow_dispatch += qty(group.summary.outflow_dispatch)
            grand_totals.outflow_transfers_out += qty(group.summary.outflow_transfers_out)
            grand_totals.total_boxes += group.summary.total_boxes
            grand_totals.total_pieces += group.summary.total_pieces
            grand_totals.closing_balance += qty(group.summary.closing_balance)

        return ProductInventoryReportResponse(
            report_type=payload.report_type,
            export_format=payload.export_format,
            generated_at=datetime.utcnow(),
            start_date=payload.start_date,
            end_date=payload.end_date,
            store=payload.store,
            product_category_id=payload.product_category_id,
            product_id=payload.product_id,
            balance_unit=payload.balance_unit,
            groups=groups,
            grand_totals=grand_totals,
        )

    def export_product_inventory_report(
        self,
        payload: ProductInventoryReportRequest,
        *,
        prepared_by: str | None = None,
    ) -> GeneratedReportFile:
        report = self.build_product_inventory_report(payload)
        export_format = payload.export_format or "pdf"
        filename_base = self._product_report_filename_base(report.generated_at)

        if export_format == "csv":
            return GeneratedReportFile(
                filename=f"{filename_base}.csv",
                content_type="text/csv",
                data=self._render_product_report_csv(report, prepared_by=prepared_by),
            )
        if export_format == "docx":
            return GeneratedReportFile(
                filename=f"{filename_base}.docx",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                data=self._render_product_report_docx(report, prepared_by=prepared_by),
            )
        return GeneratedReportFile(
            filename=f"{filename_base}.pdf",
            content_type="application/pdf",
            data=self._render_product_report_pdf(report, prepared_by=prepared_by),
        )

    # =========================================================================
    # CONSUMABLE REPORTS / EXPORT
    # =========================================================================

    def build_consumable_inventory_report(
        self,
        payload: ConsumableInventoryReportRequest,
    ) -> ConsumableInventoryReportResponse:
        query = self.db.query(ConsumableStoreInventory).filter(
            ConsumableStoreInventory.entry_date >= payload.start_date,
            ConsumableStoreInventory.entry_date <= payload.end_date,
        )

        if payload.store:
            query = query.filter(ConsumableStoreInventory.store == payload.store)
        if payload.item_category_id:
            query = query.filter(ConsumableStoreInventory.item_category_id == payload.item_category_id)
        if payload.item_id:
            query = query.filter(ConsumableStoreInventory.item_id == payload.item_id)
        if payload.unit:
            query = query.filter(ConsumableStoreInventory.unit == payload.unit)
        if payload.checked_by_initials:
            query = query.filter(ConsumableStoreInventory.checked_by_initials == payload.checked_by_initials)

        rows = query.order_by(
            ConsumableStoreInventory.entry_date.asc(),
            ConsumableStoreInventory.store.asc(),
            ConsumableStoreInventory.item_category_name.asc(),
            ConsumableStoreInventory.item_name.asc(),
            ConsumableStoreInventory.serial_no.asc(),
            ConsumableStoreInventory.created_at.asc(),
        ).all()

        groups = self._group_consumable_report_rows(rows, payload.report_type)
        grand_totals = ConsumableInventoryReportSummary()

        for group in groups:
            grand_totals.opening_balance += qty(group.summary.opening_balance)
            grand_totals.issued_today += qty(group.summary.issued_today)
            grand_totals.closing_balance += qty(group.summary.closing_balance)

        return ConsumableInventoryReportResponse(
            report_type=payload.report_type,
            export_format=payload.export_format,
            generated_at=datetime.utcnow(),
            start_date=payload.start_date,
            end_date=payload.end_date,
            store=payload.store,
            item_category_id=payload.item_category_id,
            item_id=payload.item_id,
            unit=payload.unit,
            groups=groups,
            grand_totals=grand_totals,
        )

    def export_consumable_inventory_report(
        self,
        payload: ConsumableInventoryReportRequest,
        *,
        prepared_by: str | None = None,
    ) -> GeneratedReportFile:
        report = self.build_consumable_inventory_report(payload)
        export_format = payload.export_format or "pdf"
        filename_base = self._consumable_report_filename_base(report.generated_at)

        if export_format == "csv":
            return GeneratedReportFile(
                filename=f"{filename_base}.csv",
                content_type="text/csv",
                data=self._render_consumable_report_csv(report, prepared_by=prepared_by),
            )
        if export_format == "docx":
            return GeneratedReportFile(
                filename=f"{filename_base}.docx",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                data=self._render_consumable_report_docx(report, prepared_by=prepared_by),
            )
        return GeneratedReportFile(
            filename=f"{filename_base}.pdf",
            content_type="application/pdf",
            data=self._render_consumable_report_pdf(report, prepared_by=prepared_by),
        )

    # =========================================================================
    # PRIVATE STORE NORMALIZERS
    # =========================================================================

    def _normalize_product_store(self, store: ProductStoreName | str) -> ProductStoreName:
        if isinstance(store, ProductStoreName):
            return store
        try:
            return ProductStoreName(store)
        except Exception as exc:
            raise InventoryServiceError(f"Invalid product store: {store}") from exc

    def _normalize_consumable_store(self, store: ConsumableStoreName | str) -> ConsumableStoreName:
        if isinstance(store, ConsumableStoreName):
            return store
        try:
            return ConsumableStoreName(store)
        except Exception as exc:
            raise InventoryServiceError(f"Invalid consumable store: {store}") from exc

    # =========================================================================
    # PRIVATE MODEL LIST HELPERS
    # =========================================================================

    def _list_product_models(self, active_only: bool = False) -> list[InventoryProduct]:
        query = (
            self.db.query(InventoryProduct)
            .options(joinedload(InventoryProduct.category))
            .order_by(
                InventoryProduct.sort_order.asc(),
                InventoryProduct.name.asc(),
            )
        )
        if active_only:
            query = query.filter(InventoryProduct.is_active.is_(True))
        return query.all()

    def _list_consumable_item_models(self, active_only: bool = False) -> list[InventoryConsumableItem]:
        query = (
            self.db.query(InventoryConsumableItem)
            .options(joinedload(InventoryConsumableItem.category))
            .order_by(
                InventoryConsumableItem.sort_order.asc(),
                InventoryConsumableItem.name.asc(),
            )
        )
        if active_only:
            query = query.filter(InventoryConsumableItem.is_active.is_(True))
        return query.all()

    # =========================================================================
    # PRIVATE DAILY SHEET QUERY HELPERS
    # =========================================================================

    def _list_product_entries_for_sheet(
        self,
        *,
        entry_date: date,
        store: ProductStoreName,
    ) -> list[ProductStoreInventory]:
        return (
            self.db.query(ProductStoreInventory)
            .filter(
                ProductStoreInventory.entry_date == entry_date,
                ProductStoreInventory.store == store,
            )
            .order_by(
                ProductStoreInventory.serial_no.asc(),
                ProductStoreInventory.product_category_name.asc(),
                ProductStoreInventory.product_name.asc(),
            )
            .all()
        )

    def _list_consumable_entries_for_sheet(
        self,
        *,
        entry_date: date,
        store: ConsumableStoreName,
    ) -> list[ConsumableStoreInventory]:
        return (
            self.db.query(ConsumableStoreInventory)
            .filter(
                ConsumableStoreInventory.entry_date == entry_date,
                ConsumableStoreInventory.store == store,
            )
            .order_by(
                ConsumableStoreInventory.serial_no.asc(),
                ConsumableStoreInventory.item_category_name.asc(),
                ConsumableStoreInventory.item_name.asc(),
            )
            .all()
        )

    def _product_sheet_row_from_existing(self, row: ProductStoreInventory) -> dict[str, Any]:
        read = self._to_product_inventory_read(row)
        data = read.model_dump(mode="python")
        data["entry_id"] = read.id
        data["is_existing"] = True
        data["source"] = "saved"
        return data

    def _consumable_sheet_row_from_existing(self, row: ConsumableStoreInventory) -> dict[str, Any]:
        read = self._to_consumable_inventory_read(row)
        data = read.model_dump(mode="python")
        data["entry_id"] = read.id
        data["is_existing"] = True
        data["source"] = "saved"
        return data

    def _product_sheet_row_from_product(
        self,
        *,
        product: InventoryProduct,
        entry_date: date,
        store: ProductStoreName,
        serial_no: int,
    ) -> dict[str, Any]:
        opening_balance = self.get_product_opening_balance_preview(
            entry_date=entry_date,
            store=store,
            product_id=product.id,
        )
        balance_unit = product.default_stock_unit or ProductBalanceUnit.KG
        category_name = product.category.name if product.category else ""

        return {
            "id": None,
            "entry_id": None,
            "serial_no": serial_no,
            "entry_date": entry_date,
            "store": store,
            "product_category_id": product.category_id,
            "product_id": product.id,
            "product_category_name": category_name,
            "product_name": product.name,
            "balance_unit": balance_unit,
            "opening_balance": opening_balance,
            "inflow_production": qty(0),
            "inflow_transfers_in": qty(0),
            "outflow_dispatch": qty(0),
            "outflow_transfers_out": qty(0),
            "total_boxes": 0,
            "total_pieces": 0,
            "closing_balance": opening_balance,
            "remarks": None,
            "checked_by_initials": None,
            "created_by": None,
            "updated_by": None,
            "created_at": None,
            "updated_at": None,
            "week_start_date": None,
            "week_end_date": None,
            "net_movement": qty(0),
            "is_existing": False,
            "source": "generated",
        }

    def _consumable_sheet_row_from_item(
        self,
        *,
        item: InventoryConsumableItem,
        entry_date: date,
        store: ConsumableStoreName,
        serial_no: int,
    ) -> dict[str, Any]:
        opening_balance = self.get_consumable_opening_balance_preview(
            entry_date=entry_date,
            store=store,
            item_id=item.id,
        )
        category_name = item.category.name if item.category else ""

        return {
            "id": None,
            "entry_id": None,
            "serial_no": serial_no,
            "entry_date": entry_date,
            "store": store,
            "item_category_id": item.category_id,
            "item_id": item.id,
            "item_category_name": category_name,
            "item_name": item.name,
            "unit": item.default_unit,
            "opening_balance": opening_balance,
            "issued_today": qty(0),
            "closing_balance": opening_balance,
            "remarks": None,
            "checked_by_initials": None,
            "created_by": None,
            "updated_by": None,
            "created_at": None,
            "updated_at": None,
            "week_start_date": None,
            "week_end_date": None,
            "is_existing": False,
            "source": "generated",
        }

    # =========================================================================
    # PRIVATE BULK PAYLOAD COERCION / DUPLICATE CHECKS
    # =========================================================================

    def _validate_no_duplicate_product_rows(
        self,
        rows: list[dict[str, Any] | ProductStoreInventoryCreate],
    ) -> None:
        seen: dict[str, int] = {}
        for index, raw_row in enumerate(rows, start=1):
            if isinstance(raw_row, ProductStoreInventoryCreate):
                product_id = raw_row.product_id
            elif isinstance(raw_row, dict):
                product_id = safe_text(raw_row.get("product_id")).strip()
            else:
                raise InventoryServiceError(f"Unsupported product row payload at row {index}.")

            if not product_id:
                raise InventoryServiceError(f"Product row {index}: product_id is required.")

            if product_id in seen:
                raise InventoryServiceError(
                    f"Duplicate product row in request: product_id '{product_id}' appears in rows {seen[product_id]} and {index}."
                )
            seen[product_id] = index

    def _validate_no_duplicate_consumable_rows(
        self,
        rows: list[dict[str, Any] | ConsumableStoreInventoryCreate],
    ) -> None:
        seen: dict[str, int] = {}
        for index, raw_row in enumerate(rows, start=1):
            if isinstance(raw_row, ConsumableStoreInventoryCreate):
                item_id = raw_row.item_id
            elif isinstance(raw_row, dict):
                item_id = safe_text(raw_row.get("item_id")).strip()
            else:
                raise InventoryServiceError(f"Unsupported consumable row payload at row {index}.")

            if not item_id:
                raise InventoryServiceError(f"Consumable row {index}: item_id is required.")

            if item_id in seen:
                raise InventoryServiceError(
                    f"Duplicate consumable row in request: item_id '{item_id}' appears in rows {seen[item_id]} and {index}."
                )
            seen[item_id] = index

    def _coerce_product_bulk_payload(
        self,
        *,
        raw_row: dict[str, Any] | ProductStoreInventoryCreate,
        entry_date: date,
        store: ProductStoreName,
        default_serial_no: int,
    ) -> tuple[ProductStoreInventoryCreate, set[str]]:
        if isinstance(raw_row, ProductStoreInventoryCreate):
            provided_fields = set(raw_row.model_fields_set)
            data = raw_row.model_dump(mode="python")
        elif isinstance(raw_row, dict):
            provided_fields = set(raw_row.keys())
            data = dict(raw_row)
        else:
            raise InventoryServiceError("Unsupported product bulk row payload.")

        data["entry_date"] = entry_date
        data["store"] = store
        data.setdefault("serial_no", default_serial_no)
        payload = ProductStoreInventoryCreate.model_validate(data)
        return payload, provided_fields

    def _coerce_consumable_bulk_payload(
        self,
        *,
        raw_row: dict[str, Any] | ConsumableStoreInventoryCreate,
        entry_date: date,
        store: ConsumableStoreName,
        default_serial_no: int,
    ) -> tuple[ConsumableStoreInventoryCreate, set[str]]:
        if isinstance(raw_row, ConsumableStoreInventoryCreate):
            provided_fields = set(raw_row.model_fields_set)
            data = raw_row.model_dump(mode="python")
        elif isinstance(raw_row, dict):
            provided_fields = set(raw_row.keys())
            data = dict(raw_row)
        else:
            raise InventoryServiceError("Unsupported consumable bulk row payload.")

        data["entry_date"] = entry_date
        data["store"] = store
        data.setdefault("serial_no", default_serial_no)
        payload = ConsumableStoreInventoryCreate.model_validate(data)
        return payload, provided_fields

    def _apply_product_bulk_payload_to_existing_row(
        self,
        *,
        row: ProductStoreInventory,
        payload: ProductStoreInventoryCreate,
        provided_fields: set[str],
        actor_id: str | None,
        fallback_serial_no: int,
    ) -> None:
        row.entry_date = payload.entry_date
        row.store = payload.store
        row.product_id = payload.product_id

        if payload.serial_no is not None or row.serial_no is None:
            row.serial_no = payload.serial_no or row.serial_no or fallback_serial_no

        row.product_category_id = payload.product_category_id or row.product_category_id
        row.balance_unit = payload.balance_unit or row.balance_unit

        row.inflow_production = payload.inflow_production
        row.inflow_transfers_in = payload.inflow_transfers_in
        row.outflow_dispatch = payload.outflow_dispatch
        row.outflow_transfers_out = payload.outflow_transfers_out
        row.total_boxes = payload.total_boxes
        row.total_pieces = payload.total_pieces
        row.remarks = payload.remarks
        row.checked_by_initials = payload.checked_by_initials

        if "opening_balance" in provided_fields or payload.overwrite_opening_balance:
            row.opening_balance = payload.opening_balance

        row.updated_by = actor_id or row.updated_by

    def _apply_consumable_bulk_payload_to_existing_row(
        self,
        *,
        row: ConsumableStoreInventory,
        payload: ConsumableStoreInventoryCreate,
        provided_fields: set[str],
        actor_id: str | None,
        fallback_serial_no: int,
    ) -> None:
        row.entry_date = payload.entry_date
        row.store = payload.store
        row.item_id = payload.item_id

        if payload.serial_no is not None or row.serial_no is None:
            row.serial_no = payload.serial_no or row.serial_no or fallback_serial_no

        row.item_category_id = payload.item_category_id or row.item_category_id
        row.unit = payload.unit or row.unit

        row.issued_today = payload.issued_today
        row.remarks = payload.remarks
        row.checked_by_initials = payload.checked_by_initials

        if "opening_balance" in provided_fields or payload.overwrite_opening_balance:
            row.opening_balance = payload.opening_balance

        row.updated_by = actor_id or row.updated_by

    # =========================================================================
    # PRIVATE INVENTORY BUILD / PREP HELPERS
    # =========================================================================

    def _build_new_product_inventory_row(
        self,
        payload: ProductStoreInventoryCreate,
    ) -> ProductStoreInventory:
        data = payload.model_dump(exclude={"overwrite_opening_balance"}, exclude_none=True)
        return ProductStoreInventory(**data)

    def _build_new_consumable_inventory_row(
        self,
        payload: ConsumableStoreInventoryCreate,
    ) -> ConsumableStoreInventory:
        data = payload.model_dump(exclude={"overwrite_opening_balance"}, exclude_none=True)
        return ConsumableStoreInventory(**data)

    def _prepare_product_inventory_row(
        self,
        row: ProductStoreInventory,
        *,
        overwrite_opening_balance: bool,
        exclude_entry_id: str | None = None,
    ) -> None:
        product = self.get_product_or_404(row.product_id)

        if not row.product_category_id:
            row.product_category_id = product.category_id
        if not row.balance_unit:
            row.balance_unit = product.default_stock_unit or ProductBalanceUnit.KG

        self._apply_product_opening_balance_rule(
            row,
            overwrite_opening_balance=overwrite_opening_balance,
            exclude_entry_id=exclude_entry_id,
        )
        row.prepare_for_save(self.db, overwrite_opening_balance=False)
        self._ensure_product_stock_is_non_negative(row)

    def _prepare_consumable_inventory_row(
        self,
        row: ConsumableStoreInventory,
        *,
        overwrite_opening_balance: bool,
        exclude_entry_id: str | None = None,
    ) -> None:
        item = self.get_consumable_item_or_404(row.item_id)

        if not row.item_category_id:
            row.item_category_id = item.category_id
        if not row.unit:
            row.unit = item.default_unit

        self._apply_consumable_opening_balance_rule(
            row,
            overwrite_opening_balance=overwrite_opening_balance,
            exclude_entry_id=exclude_entry_id,
        )
        row.prepare_for_save(self.db, overwrite_opening_balance=False)
        self._ensure_consumable_stock_is_non_negative(row)

    # =========================================================================
    # PRIVATE VALIDATION / OPENING BALANCE RULES
    # =========================================================================

    def _get_existing_product_entry(
        self,
        *,
        entry_date: date,
        store: ProductStoreName | str,
        product_id: str,
    ) -> ProductStoreInventory | None:
        return (
            self.db.query(ProductStoreInventory)
            .filter(
                ProductStoreInventory.entry_date == entry_date,
                ProductStoreInventory.store == store,
                ProductStoreInventory.product_id == product_id,
            )
            .first()
        )

    def _get_existing_consumable_entry(
        self,
        *,
        entry_date: date,
        store: ConsumableStoreName | str,
        item_id: str,
    ) -> ConsumableStoreInventory | None:
        return (
            self.db.query(ConsumableStoreInventory)
            .filter(
                ConsumableStoreInventory.entry_date == entry_date,
                ConsumableStoreInventory.store == store,
                ConsumableStoreInventory.item_id == item_id,
            )
            .first()
        )

    def _get_previous_product_entry(
        self,
        *,
        entry_date: date,
        store: ProductStoreName | str,
        product_id: str,
        exclude_entry_id: str | None = None,
    ) -> ProductStoreInventory | None:
        query = self.db.query(ProductStoreInventory).filter(
            ProductStoreInventory.entry_date < entry_date,
            ProductStoreInventory.store == store,
            ProductStoreInventory.product_id == product_id,
        )
        if exclude_entry_id:
            query = query.filter(ProductStoreInventory.id != exclude_entry_id)
        return query.order_by(
            ProductStoreInventory.entry_date.desc(),
            ProductStoreInventory.created_at.desc(),
        ).first()

    def _get_previous_consumable_entry(
        self,
        *,
        entry_date: date,
        store: ConsumableStoreName | str,
        item_id: str,
        exclude_entry_id: str | None = None,
    ) -> ConsumableStoreInventory | None:
        query = self.db.query(ConsumableStoreInventory).filter(
            ConsumableStoreInventory.entry_date < entry_date,
            ConsumableStoreInventory.store == store,
            ConsumableStoreInventory.item_id == item_id,
        )
        if exclude_entry_id:
            query = query.filter(ConsumableStoreInventory.id != exclude_entry_id)
        return query.order_by(
            ConsumableStoreInventory.entry_date.desc(),
            ConsumableStoreInventory.created_at.desc(),
        ).first()

    def _apply_product_opening_balance_rule(
        self,
        row: ProductStoreInventory,
        *,
        overwrite_opening_balance: bool,
        exclude_entry_id: str | None = None,
    ) -> None:
        previous = self._get_previous_product_entry(
            entry_date=row.entry_date,
            store=row.store,
            product_id=row.product_id,
            exclude_entry_id=exclude_entry_id,
        )
        expected_opening = qty(previous.closing_balance if previous else 0)

        if row.opening_balance is None or overwrite_opening_balance:
            row.opening_balance = expected_opening
            return

        supplied_opening = qty(row.opening_balance)
        if previous and not decimal_equal(supplied_opening, expected_opening):
            raise InventoryServiceError(
                f"Opening balance must match previous closing balance ({expected_opening}). "
                "Leave opening_balance blank for auto-fill or set overwrite_opening_balance=true."
            )

        row.opening_balance = supplied_opening

    def _apply_consumable_opening_balance_rule(
        self,
        row: ConsumableStoreInventory,
        *,
        overwrite_opening_balance: bool,
        exclude_entry_id: str | None = None,
    ) -> None:
        previous = self._get_previous_consumable_entry(
            entry_date=row.entry_date,
            store=row.store,
            item_id=row.item_id,
            exclude_entry_id=exclude_entry_id,
        )
        expected_opening = qty(previous.closing_balance if previous else 0)

        if row.opening_balance is None or overwrite_opening_balance:
            row.opening_balance = expected_opening
            return

        supplied_opening = qty(row.opening_balance)
        if previous and not decimal_equal(supplied_opening, expected_opening):
            raise InventoryServiceError(
                f"Opening balance must match previous closing balance ({expected_opening}). "
                "Leave opening_balance blank for auto-fill or set overwrite_opening_balance=true."
            )

        row.opening_balance = supplied_opening

    def _ensure_product_stock_is_non_negative(self, row: ProductStoreInventory) -> None:
        available_before_outflow = (
            qty(row.opening_balance)
            + qty(row.inflow_production)
            + qty(row.inflow_transfers_in)
        )
        total_outflow = qty(row.outflow_dispatch) + qty(row.outflow_transfers_out)

        if total_outflow > available_before_outflow:
            raise InventoryServiceError(
                "Product stock cannot go negative. "
                f"Available stock is {available_before_outflow}, but total outflow is {total_outflow}."
            )

        if qty(row.closing_balance) < 0:
            raise InventoryServiceError(
                f"Closing balance cannot be negative. Computed closing balance is {qty(row.closing_balance)}."
            )

    def _ensure_consumable_stock_is_non_negative(self, row: ConsumableStoreInventory) -> None:
        available_before_issue = qty(row.opening_balance)
        issued_today = qty(row.issued_today)

        if issued_today > available_before_issue:
            raise InventoryServiceError(
                "Consumable stock cannot go negative. "
                f"Opening balance is {available_before_issue}, but issued today is {issued_today}."
            )

        if qty(row.closing_balance) < 0:
            raise InventoryServiceError(
                f"Closing balance cannot be negative. Computed closing balance is {qty(row.closing_balance)}."
            )

    # =========================================================================
    # PRIVATE READ MAPPERS
    # =========================================================================

    def _coerce_product_store_enum(self, value: Any) -> ProductStoreName | None:
        if value is None:
            return None
        if isinstance(value, ProductStoreName):
            return value

        if isinstance(value, str):
            raw = value.strip()
            if not raw:
                return None

            # Prefer direct enum value match, e.g. "Chiller 1"
            try:
                return ProductStoreName(raw)
            except Exception:
                pass

            # Fallback to enum member name, e.g. "CHILLER_1"
            try:
                return ProductStoreName[raw]
            except Exception:
                pass

            normalized = raw.upper().replace(" ", "_").replace("-", "_")
            try:
                return ProductStoreName[normalized]
            except Exception:
                pass

        raise InventoryServiceError(f"Invalid product store value: {value!r}")

    def _coerce_consumable_store_enum(self, value: Any) -> ConsumableStoreName | None:
        if value is None:
            return None
        if isinstance(value, ConsumableStoreName):
            return value

        if isinstance(value, str):
            raw = value.strip()
            if not raw:
                return None

            try:
                return ConsumableStoreName(raw)
            except Exception:
                pass

            try:
                return ConsumableStoreName[raw]
            except Exception:
                pass

            normalized = raw.upper().replace(" ", "_").replace("-", "_")
            try:
                return ConsumableStoreName[normalized]
            except Exception:
                pass

        raise InventoryServiceError(f"Invalid consumable store value: {value!r}")

    def _coerce_balance_unit_enum(self, value: Any) -> ProductBalanceUnit | None:
        if value is None:
            return None
        if isinstance(value, ProductBalanceUnit):
            return value

        if isinstance(value, str):
            raw = value.strip()
            if not raw:
                return None

            lowered = raw.lower()
            if lowered in {"kg", "kgs"}:
                return ProductBalanceUnit.KG
            if lowered in {"pcs", "pc", "piece", "pieces"}:
                return ProductBalanceUnit.PCS

            try:
                return ProductBalanceUnit(raw)
            except Exception:
                pass

            try:
                return ProductBalanceUnit[raw]
            except Exception:
                pass

            normalized = raw.upper().replace(" ", "_").replace("-", "_")
            try:
                return ProductBalanceUnit[normalized]
            except Exception:
                pass

        raise InventoryServiceError(f"Invalid product balance unit value: {value!r}")

    def _to_product_inventory_read(self, row: ProductStoreInventory) -> ProductStoreInventoryRead:
        data = {
            "id": row.id,
            "serial_no": row.serial_no,
            "entry_date": row.entry_date,
            "store": self._coerce_product_store_enum(getattr(row, "store", None)),
            "product_category_id": row.product_category_id,
            "product_id": row.product_id,
            "product_category_name": row.product_category_name,
            "product_name": row.product_name,
            "balance_unit": self._coerce_balance_unit_enum(getattr(row, "balance_unit", None)),
            "opening_balance": qty(row.opening_balance),
            "inflow_production": qty(row.inflow_production),
            "inflow_transfers_in": qty(row.inflow_transfers_in),
            "outflow_dispatch": qty(row.outflow_dispatch),
            "outflow_transfers_out": qty(row.outflow_transfers_out),
            "total_boxes": int(row.total_boxes or 0),
            "total_pieces": int(row.total_pieces or 0),
            "closing_balance": qty(row.closing_balance),
            "remarks": row.remarks,
            "checked_by_initials": row.checked_by_initials,
            "created_by": row.created_by,
            "updated_by": row.updated_by,
            "created_at": row.created_at,
            "updated_at": row.updated_at,
            "week_start_date": getattr(row, "week_start_date", None),
            "week_end_date": getattr(row, "week_end_date", None),
            "net_movement": qty(getattr(row, "net_movement", None)),
        }
        return ProductStoreInventoryRead.model_validate(data)

    def _to_consumable_inventory_read(self, row: ConsumableStoreInventory) -> ConsumableStoreInventoryRead:
        data = {
            "id": row.id,
            "serial_no": row.serial_no,
            "entry_date": row.entry_date,
            "store": self._coerce_consumable_store_enum(getattr(row, "store", None)),
            "item_category_id": row.item_category_id,
            "item_id": row.item_id,
            "item_category_name": row.item_category_name,
            "item_name": row.item_name,
            "unit": safe_text(row.unit),
            "opening_balance": qty(row.opening_balance),
            "issued_today": qty(row.issued_today),
            "closing_balance": qty(row.closing_balance),
            "remarks": row.remarks,
            "checked_by_initials": row.checked_by_initials,
            "created_by": row.created_by,
            "updated_by": row.updated_by,
            "created_at": row.created_at,
            "updated_at": row.updated_at,
            "week_start_date": getattr(row, "week_start_date", None),
            "week_end_date": getattr(row, "week_end_date", None),
        }
        return ConsumableStoreInventoryRead.model_validate(data)

    # =========================================================================
    # PRIVATE COMMIT / CONFLICT
    # =========================================================================

    def _commit_or_conflict(self, message: str) -> None:
        try:
            self.db.commit()
        except IntegrityError as exc:
            self.db.rollback()
            raise InventoryConflictError(message) from exc

    # =========================================================================
    # PRIVATE SERIAL NUMBER HELPERS
    # =========================================================================

    def _next_product_serial_no(self, entry_date: date, store: ProductStoreName | str) -> int:
        current_max = (
            self.db.query(func.max(ProductStoreInventory.serial_no))
            .filter(
                ProductStoreInventory.entry_date == entry_date,
                ProductStoreInventory.store == store,
            )
            .scalar()
        )
        return int(current_max or 0) + 1

    def _next_consumable_serial_no(self, entry_date: date, store: ConsumableStoreName | str) -> int:
        current_max = (
            self.db.query(func.max(ConsumableStoreInventory.serial_no))
            .filter(
                ConsumableStoreInventory.entry_date == entry_date,
                ConsumableStoreInventory.store == store,
            )
            .scalar()
        )
        return int(current_max or 0) + 1

    # =========================================================================
    # PRIVATE REPORT GROUPING (FIXED)
    # =========================================================================

    def _group_product_report_rows(
        self,
        rows: list[ProductStoreInventory],
        report_type: str,
    ) -> list[ProductInventoryReportGroup]:
        grouped: dict[tuple[date, date, str], list[ProductStoreInventory]] = defaultdict(list)

        for row in rows:
            if report_type == "weekly":
                start = row.week_start_date
                end = row.week_end_date
                label = f"Week: {fmt_date(start)} to {fmt_date(end)}"
            else:
                start = row.entry_date
                end = row.entry_date
                label = f"Date: {fmt_date(row.entry_date)}"

            grouped[(start, end, label)].append(row)

        groups: list[ProductInventoryReportGroup] = []

        for (start, end, label), items in sorted(grouped.items(), key=lambda x: x[0][0]):
            summary = ProductInventoryReportSummary()
            report_rows: list[ProductInventoryReportRow] = []

            # FIX: Use enumeration index for serial_no, ignore stored value
            for index, row in enumerate(items, start=1):
                report_rows.append(
                    ProductInventoryReportRow(
                        serial_no=index,    # <-- sequential index
                        entry_date=row.entry_date,
                        store=row.store,
                        product_category_name=row.product_category_name,
                        product_name=row.product_name,
                        balance_unit=row.balance_unit,
                        opening_balance=qty(row.opening_balance),
                        inflow_production=qty(row.inflow_production),
                        inflow_transfers_in=qty(row.inflow_transfers_in),
                        outflow_dispatch=qty(row.outflow_dispatch),
                        outflow_transfers_out=qty(row.outflow_transfers_out),
                        total_boxes=int(row.total_boxes or 0),
                        total_pieces=int(row.total_pieces or 0),
                        closing_balance=qty(row.closing_balance),
                        remarks=row.remarks,
                        checked_by_initials=row.checked_by_initials,
                    )
                )

                summary.opening_balance += qty(row.opening_balance)
                summary.inflow_production += qty(row.inflow_production)
                summary.inflow_transfers_in += qty(row.inflow_transfers_in)
                summary.outflow_dispatch += qty(row.outflow_dispatch)
                summary.outflow_transfers_out += qty(row.outflow_transfers_out)
                summary.total_boxes += int(row.total_boxes or 0)
                summary.total_pieces += int(row.total_pieces or 0)
                summary.closing_balance += qty(row.closing_balance)

            groups.append(
                ProductInventoryReportGroup(
                    label=label,
                    start_date=start,
                    end_date=end,
                    rows=report_rows,
                    summary=summary,
                )
            )

        return groups

    def _group_consumable_report_rows(
        self,
        rows: list[ConsumableStoreInventory],
        report_type: str,
    ) -> list[ConsumableInventoryReportGroup]:
        grouped: dict[tuple[date, date, str], list[ConsumableStoreInventory]] = defaultdict(list)

        for row in rows:
            if report_type == "weekly":
                start = row.week_start_date
                end = row.week_end_date
                label = f"Week: {fmt_date(start)} to {fmt_date(end)}"
            else:
                start = row.entry_date
                end = row.entry_date
                label = f"Date: {fmt_date(row.entry_date)}"

            grouped[(start, end, label)].append(row)

        groups: list[ConsumableInventoryReportGroup] = []

        for (start, end, label), items in sorted(grouped.items(), key=lambda x: x[0][0]):
            summary = ConsumableInventoryReportSummary()
            report_rows: list[ConsumableInventoryReportRow] = []

            # FIX: Use enumeration index for serial_no, ignore stored value
            for index, row in enumerate(items, start=1):
                report_rows.append(
                    ConsumableInventoryReportRow(
                        serial_no=index,    # <-- sequential index
                        entry_date=row.entry_date,
                        store=row.store,
                        item_category_name=row.item_category_name,
                        item_name=row.item_name,
                        unit=row.unit,
                        opening_balance=qty(row.opening_balance),
                        issued_today=qty(row.issued_today),
                        closing_balance=qty(row.closing_balance),
                        remarks=row.remarks,
                        checked_by_initials=row.checked_by_initials,
                    )
                )

                summary.opening_balance += qty(row.opening_balance)
                summary.issued_today += qty(row.issued_today)
                summary.closing_balance += qty(row.closing_balance)

            groups.append(
                ConsumableInventoryReportGroup(
                    label=label,
                    start_date=start,
                    end_date=end,
                    rows=report_rows,
                    summary=summary,
                )
            )

        return groups

    # =========================================================================
    # PRIVATE REPORT HELPERS
    # =========================================================================

    def _resolve_prepared_by(self, prepared_by: str | None) -> str:
        value = safe_text(prepared_by).strip()
        return value or "Logged in User"

    def _summary_label(self) -> str:
        return "Grand Total"

    def _product_report_filename_base(self, generated_at: datetime) -> str:
        return f"Product Inventory Report {generated_at.strftime('%Y%m%d')}"

    def _consumable_report_filename_base(self, generated_at: datetime) -> str:
        return f"Consumable Inventory Report {generated_at.strftime('%Y%m%d')}"

    # =========================================================================
    # CSV RENDERING
    # =========================================================================

    def _render_product_report_csv(
        self,
        report: ProductInventoryReportResponse,
        *,
        prepared_by: str | None = None,
    ) -> bytes:
        prepared_by = self._resolve_prepared_by(prepared_by)
        buffer = StringIO()
        writer = csv.writer(buffer)

        writer.writerow(["PRODUCT STORE INVENTORY REPORT"])
        writer.writerow(["Prepared By", prepared_by])
        writer.writerow(["Report Type", report.report_type.title()])
        writer.writerow(["Period", f"{fmt_date(report.start_date)} to {fmt_date(report.end_date)}"])
        writer.writerow(["Generated At", report.generated_at.isoformat()])
        writer.writerow([])

        for group in report.groups:
            writer.writerow([group.label])
            writer.writerow(
                [
                    "S.No.",
                    "Date",
                    "Store",
                    "Product Category",
                    "Product Type",
                    "Opening Balance",
                    "Inflow - Production",
                    "Inflow - Transfers In",
                    "Outflow - Dispatch",
                    "Outflow - Transfers Out",
                    "Total Boxes",
                    "Total Pieces",
                    "Closing Balance",
                    "Remarks",
                    "Checked By",
                ]
            )

            for row in group.rows:
                writer.writerow(
                    [
                        row.serial_no,
                        fmt_date(row.entry_date),
                        row.store.value,
                        row.product_category_name,
                        row.product_name,
                        fmt_qty(row.opening_balance),
                        fmt_qty(row.inflow_production),
                        fmt_qty(row.inflow_transfers_in),
                        fmt_qty(row.outflow_dispatch),
                        fmt_qty(row.outflow_transfers_out),
                        row.total_boxes,
                        row.total_pieces,
                        fmt_qty(row.closing_balance),
                        safe_text(row.remarks),
                        safe_text(row.checked_by_initials),
                    ]
                )

            writer.writerow(
                [
                    "",
                    "",
                    "",
                    "",
                    self._summary_label(),
                    fmt_qty(group.summary.opening_balance),
                    fmt_qty(group.summary.inflow_production),
                    fmt_qty(group.summary.inflow_transfers_in),
                    fmt_qty(group.summary.outflow_dispatch),
                    fmt_qty(group.summary.outflow_transfers_out),
                    group.summary.total_boxes,
                    group.summary.total_pieces,
                    fmt_qty(group.summary.closing_balance),
                    "",
                    "",
                ]
            )
            writer.writerow([])

        return buffer.getvalue().encode("utf-8")

    def _render_consumable_report_csv(
        self,
        report: ConsumableInventoryReportResponse,
        *,
        prepared_by: str | None = None,
    ) -> bytes:
        prepared_by = self._resolve_prepared_by(prepared_by)
        buffer = StringIO()
        writer = csv.writer(buffer)

        writer.writerow(["CONSUMABLE STORE INVENTORY REPORT"])
        writer.writerow(["Prepared By", prepared_by])
        writer.writerow(["Report Type", report.report_type.title()])
        writer.writerow(["Period", f"{fmt_date(report.start_date)} to {fmt_date(report.end_date)}"])
        writer.writerow(["Generated At", report.generated_at.isoformat()])
        writer.writerow([])

        for group in report.groups:
            writer.writerow([group.label])
            writer.writerow(
                [
                    "S.No.",
                    "Date",
                    "Store",
                    "Item Category",
                    "Item",
                    "Unit",
                    "Opening Balance",
                    "Issued Today",
                    "Closing Balance",
                    "Remarks",
                    "Checked By",
                ]
            )

            for row in group.rows:
                writer.writerow(
                    [
                        row.serial_no,
                        fmt_date(row.entry_date),
                        row.store.value,
                        row.item_category_name,
                        row.item_name,
                        row.unit,
                        fmt_qty(row.opening_balance),
                        fmt_qty(row.issued_today),
                        fmt_qty(row.closing_balance),
                        safe_text(row.remarks),
                        safe_text(row.checked_by_initials),
                    ]
                )

            writer.writerow(
                [
                    "",
                    "",
                    "",
                    "",
                    self._summary_label(),
                    "",
                    fmt_qty(group.summary.opening_balance),
                    fmt_qty(group.summary.issued_today),
                    fmt_qty(group.summary.closing_balance),
                    "",
                    "",
                ]
            )
            writer.writerow([])

        return buffer.getvalue().encode("utf-8")

    # =========================================================================
    # PDF RENDERING
    # =========================================================================

    def _render_product_report_pdf(
        self,
        report: ProductInventoryReportResponse,
        *,
        prepared_by: str | None = None,
    ) -> bytes:
        try:
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
            from reportlab.lib.pagesizes import A4, landscape
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import mm
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
        except ImportError as exc:
            raise RuntimeError("reportlab is required to generate PDF reports.") from exc

        prepared_by = self._resolve_prepared_by(prepared_by)

        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=landscape(A4),
            leftMargin=8 * mm,
            rightMargin=8 * mm,
            topMargin=8 * mm,
            bottomMargin=8 * mm,
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "ReportTitle",
            parent=styles["Title"],
            fontName="Times-Bold",
            fontSize=18,
            leading=22,
            alignment=TA_CENTER,
            spaceAfter=8,
        )
        meta_style = ParagraphStyle(
            "ReportMeta",
            parent=styles["Normal"],
            fontName="Times-Roman",
            fontSize=11,
            leading=14,
            alignment=TA_LEFT,
            spaceAfter=2,
        )
        group_style = ParagraphStyle(
            "GroupLabel",
            parent=styles["Heading3"],
            fontName="Times-Bold",
            fontSize=14,
            leading=17,
            spaceBefore=8,
            spaceAfter=6,
        )
        header_style = ParagraphStyle(
            "HeaderCell",
            parent=styles["BodyText"],
            fontName="Times-Bold",
            fontSize=10,
            leading=12,
            alignment=TA_CENTER,
        )
        cell_left = ParagraphStyle(
            "CellLeft",
            parent=styles["BodyText"],
            fontName="Times-Roman",
            fontSize=10,
            leading=12,
            alignment=TA_LEFT,
        )
        cell_center = ParagraphStyle(
            "CellCenter",
            parent=styles["BodyText"],
            fontName="Times-Roman",
            fontSize=10,
            leading=12,
            alignment=TA_CENTER,
        )
        cell_right = ParagraphStyle(
            "CellRight",
            parent=styles["BodyText"],
            fontName="Times-Roman",
            fontSize=10,
            leading=12,
            alignment=TA_RIGHT,
        )
        total_left = ParagraphStyle(
            "TotalLeft",
            parent=cell_left,
            fontName="Times-Bold",
            fontSize=10,
            leading=12,
        )
        total_center = ParagraphStyle(
            "TotalCenter",
            parent=cell_center,
            fontName="Times-Bold",
            fontSize=10,
            leading=12,
        )
        total_right = ParagraphStyle(
            "TotalRight",
            parent=cell_right,
            fontName="Times-Bold",
            fontSize=10,
            leading=12,
        )

        story = []
        story.append(Paragraph("PRODUCT STORE INVENTORY REPORT", title_style))
        story.append(Paragraph(f"Prepared By: {safe_text(prepared_by)}", meta_style))
        story.append(Paragraph(f"Report Type: {report.report_type.title()}", meta_style))
        story.append(
            Paragraph(
                f"Period: {fmt_date(report.start_date)} to {fmt_date(report.end_date)}",
                meta_style,
            )
        )
        story.append(Paragraph(f"Generated At: {report.generated_at.isoformat()}", meta_style))
        story.append(Spacer(1, 6))

        col_widths = [
            11 * mm,
            17 * mm,
            19 * mm,
            25 * mm,
            32 * mm,
            16 * mm,
            16 * mm,
            18 * mm,
            16 * mm,
            18 * mm,
            13 * mm,
            13 * mm,
            17 * mm,
            40 * mm,
            15 * mm,
        ]

        for group in report.groups:
            story.append(Paragraph(f"{group.label}", group_style))

            table_data = [
                [
                    Paragraph("S.No.", header_style),
                    Paragraph("Date", header_style),
                    Paragraph("Store", header_style),
                    Paragraph("Category", header_style),
                    Paragraph("Product", header_style),
                    Paragraph("Opening", header_style),
                    Paragraph("Production", header_style),
                    Paragraph("Transfers<br/>In", header_style),
                    Paragraph("Dispatch", header_style),
                    Paragraph("Transfers<br/>Out", header_style),
                    Paragraph("Boxes", header_style),
                    Paragraph("Pieces", header_style),
                    Paragraph("Closing", header_style),
                    Paragraph("Remarks", header_style),
                    Paragraph("Checked", header_style),
                ]
            ]

            for row in group.rows:
                table_data.append(
                    [
                        Paragraph(safe_text(row.serial_no), cell_center),
                        Paragraph(fmt_date(row.entry_date), cell_center),
                        Paragraph(safe_text(row.store.value), cell_left),
                        Paragraph(safe_text(row.product_category_name), cell_left),
                        Paragraph(safe_text(row.product_name), cell_left),
                        Paragraph(fmt_qty(row.opening_balance), cell_right),
                        Paragraph(fmt_qty(row.inflow_production), cell_right),
                        Paragraph(fmt_qty(row.inflow_transfers_in), cell_right),
                        Paragraph(fmt_qty(row.outflow_dispatch), cell_right),
                        Paragraph(fmt_qty(row.outflow_transfers_out), cell_right),
                        Paragraph(safe_text(row.total_boxes), cell_right),
                        Paragraph(safe_text(row.total_pieces), cell_right),
                        Paragraph(fmt_qty(row.closing_balance), cell_right),
                        Paragraph(safe_text(row.remarks), cell_left),
                        Paragraph(safe_text(row.checked_by_initials), cell_center),
                    ]
                )

            table_data.append(
                [
                    Paragraph("", total_center),
                    Paragraph("", total_center),
                    Paragraph("", total_left),
                    Paragraph("", total_left),
                    Paragraph(self._summary_label(), total_left),
                    Paragraph(fmt_qty(group.summary.opening_balance), total_right),
                    Paragraph(fmt_qty(group.summary.inflow_production), total_right),
                    Paragraph(fmt_qty(group.summary.inflow_transfers_in), total_right),
                    Paragraph(fmt_qty(group.summary.outflow_dispatch), total_right),
                    Paragraph(fmt_qty(group.summary.outflow_transfers_out), total_right),
                    Paragraph(safe_text(group.summary.total_boxes), total_right),
                    Paragraph(safe_text(group.summary.total_pieces), total_right),
                    Paragraph(fmt_qty(group.summary.closing_balance), total_right),
                    Paragraph("", total_left),
                    Paragraph("", total_center),
                ]
            )

            table = Table(table_data, colWidths=col_widths, repeatRows=1, hAlign="LEFT")
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E5E7EB")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                        ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#F3F4F6")),
                        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#9CA3AF")),
                        ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#D1D5DB")),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 8))

        doc.build(story)
        return buffer.getvalue()

    def _render_consumable_report_pdf(
        self,
        report: ConsumableInventoryReportResponse,
        *,
        prepared_by: str | None = None,
    ) -> bytes:
        try:
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
            from reportlab.lib.pagesizes import A4, landscape
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import mm
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
        except ImportError as exc:
            raise RuntimeError("reportlab is required to generate PDF reports.") from exc

        prepared_by = self._resolve_prepared_by(prepared_by)

        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=landscape(A4),
            leftMargin=8 * mm,
            rightMargin=8 * mm,
            topMargin=8 * mm,
            bottomMargin=8 * mm,
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "ReportTitle",
            parent=styles["Title"],
            fontName="Times-Bold",
            fontSize=18,
            leading=22,
            alignment=TA_CENTER,
            spaceAfter=8,
        )
        meta_style = ParagraphStyle(
            "ReportMeta",
            parent=styles["Normal"],
            fontName="Times-Roman",
            fontSize=11,
            leading=14,
            alignment=TA_LEFT,
            spaceAfter=2,
        )
        group_style = ParagraphStyle(
            "GroupLabel",
            parent=styles["Heading3"],
            fontName="Times-Bold",
            fontSize=14,
            leading=17,
            spaceBefore=8,
            spaceAfter=6,
        )
        header_style = ParagraphStyle(
            "HeaderCell",
            parent=styles["BodyText"],
            fontName="Times-Bold",
            fontSize=10,
            leading=12,
            alignment=TA_CENTER,
        )
        cell_left = ParagraphStyle(
            "CellLeft",
            parent=styles["BodyText"],
            fontName="Times-Roman",
            fontSize=10,
            leading=12,
            alignment=TA_LEFT,
        )
        cell_center = ParagraphStyle(
            "CellCenter",
            parent=styles["BodyText"],
            fontName="Times-Roman",
            fontSize=10,
            leading=12,
            alignment=TA_CENTER,
        )
        cell_right = ParagraphStyle(
            "CellRight",
            parent=styles["BodyText"],
            fontName="Times-Roman",
            fontSize=10,
            leading=12,
            alignment=TA_RIGHT,
        )
        total_left = ParagraphStyle(
            "TotalLeft",
            parent=cell_left,
            fontName="Times-Bold",
            fontSize=10,
            leading=12,
        )
        total_center = ParagraphStyle(
            "TotalCenter",
            parent=cell_center,
            fontName="Times-Bold",
            fontSize=10,
            leading=12,
        )
        total_right = ParagraphStyle(
            "TotalRight",
            parent=cell_right,
            fontName="Times-Bold",
            fontSize=10,
            leading=12,
        )

        story = []
        story.append(Paragraph("CONSUMABLE STORE INVENTORY REPORT", title_style))
        story.append(Paragraph(f"Prepared By: {safe_text(prepared_by)}", meta_style))
        story.append(Paragraph(f"Report Type: {report.report_type.title()}", meta_style))
        story.append(
            Paragraph(
                f"Period: {fmt_date(report.start_date)} to {fmt_date(report.end_date)}",
                meta_style,
            )
        )
        story.append(Paragraph(f"Generated At: {report.generated_at.isoformat()}", meta_style))
        story.append(Spacer(1, 6))

        col_widths = [
            11 * mm,
            17 * mm,
            19 * mm,
            32 * mm,
            42 * mm,
            16 * mm,
            19 * mm,
            19 * mm,
            19 * mm,
            42 * mm,
            16 * mm,
        ]

        for group in report.groups:
            story.append(Paragraph(f"{group.label}", group_style))

            table_data = [
                [
                    Paragraph("S.No.", header_style),
                    Paragraph("Date", header_style),
                    Paragraph("Store", header_style),
                    Paragraph("Item<br/>Category", header_style),
                    Paragraph("Item", header_style),
                    Paragraph("Unit", header_style),
                    Paragraph("Opening", header_style),
                    Paragraph("Issued<br/>Today", header_style),
                    Paragraph("Closing", header_style),
                    Paragraph("Remarks", header_style),
                    Paragraph("Checked", header_style),
                ]
            ]

            for row in group.rows:
                table_data.append(
                    [
                        Paragraph(safe_text(row.serial_no), cell_center),
                        Paragraph(fmt_date(row.entry_date), cell_center),
                        Paragraph(safe_text(row.store.value), cell_left),
                        Paragraph(safe_text(row.item_category_name), cell_left),
                        Paragraph(safe_text(row.item_name), cell_left),
                        Paragraph(safe_text(row.unit), cell_center),
                        Paragraph(fmt_qty(row.opening_balance), cell_right),
                        Paragraph(fmt_qty(row.issued_today), cell_right),
                        Paragraph(fmt_qty(row.closing_balance), cell_right),
                        Paragraph(safe_text(row.remarks), cell_left),
                        Paragraph(safe_text(row.checked_by_initials), cell_center),
                    ]
                )

            table_data.append(
                [
                    Paragraph("", total_center),
                    Paragraph("", total_center),
                    Paragraph("", total_left),
                    Paragraph("", total_left),
                    Paragraph(self._summary_label(), total_left),
                    Paragraph("", total_center),
                    Paragraph(fmt_qty(group.summary.opening_balance), total_right),
                    Paragraph(fmt_qty(group.summary.issued_today), total_right),
                    Paragraph(fmt_qty(group.summary.closing_balance), total_right),
                    Paragraph("", total_left),
                    Paragraph("", total_center),
                ]
            )

            table = Table(table_data, colWidths=col_widths, repeatRows=1, hAlign="LEFT")
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E5E7EB")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                        ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#F3F4F6")),
                        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#9CA3AF")),
                        ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#D1D5DB")),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 8))

        doc.build(story)
        return buffer.getvalue()

    # =========================================================================
    # DOCX RENDERING
    # =========================================================================

    def _render_product_report_docx(
        self,
        report: ProductInventoryReportResponse,
        *,
        prepared_by: str | None = None,
    ) -> bytes:
        try:
            from docx import Document
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches, Pt
        except ImportError as exc:
            raise RuntimeError("python-docx is required to generate Word reports.") from exc

        prepared_by = self._resolve_prepared_by(prepared_by)
        document = Document()
        self._set_docx_landscape(document)

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("PRODUCT STORE INVENTORY REPORT")
        title_run.bold = True
        title_run.font.size = Pt(14)

        for line in [
            f"Prepared By: {prepared_by}",
            f"Report Type: {report.report_type.title()}",
            f"Period: {fmt_date(report.start_date)} to {fmt_date(report.end_date)}",
            f"Generated At: {report.generated_at.isoformat()}",
        ]:
            p = document.add_paragraph()
            r = p.add_run(line)
            r.font.size = Pt(9)

        product_widths = [
            0.36, 0.62, 0.72, 0.90, 1.00,
            0.62, 0.68, 0.74, 0.68, 0.74,
            0.46, 0.46, 0.68, 1.10, 0.56,
        ]

        for group in report.groups:
            document.add_paragraph(group.label, style="Heading 2")

            headers = [
                "S.No.",
                "Date",
                "Store",
                "Category",
                "Product",
                "Opening",
                "Production",
                "Transfers In",
                "Dispatch",
                "Transfers Out",
                "Boxes",
                "Pieces",
                "Closing",
                "Remarks",
                "Checked",
            ]

            table = document.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            table.autofit = False
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for idx, header in enumerate(headers):
                cell = table.rows[0].cells[idx]
                cell.width = Inches(product_widths[idx])
                self._write_docx_cell(
                    cell,
                    header,
                    bold=True,
                    align="center",
                    font_size=8,
                    shading="E5E7EB",
                )

            for row in group.rows:
                cells = table.add_row().cells
                values = [
                    safe_text(row.serial_no),
                    fmt_date(row.entry_date),
                    row.store.value,
                    row.product_category_name,
                    row.product_name,
                    fmt_qty(row.opening_balance),
                    fmt_qty(row.inflow_production),
                    fmt_qty(row.inflow_transfers_in),
                    fmt_qty(row.outflow_dispatch),
                    fmt_qty(row.outflow_transfers_out),
                    safe_text(row.total_boxes),
                    safe_text(row.total_pieces),
                    fmt_qty(row.closing_balance),
                    safe_text(row.remarks),
                    safe_text(row.checked_by_initials),
                ]
                aligns = [
                    "center", "center", "left", "left", "left",
                    "right", "right", "right", "right", "right",
                    "right", "right", "right", "left", "center",
                ]
                for idx, value in enumerate(values):
                    cells[idx].width = Inches(product_widths[idx])
                    self._write_docx_cell(
                        cells[idx],
                        value,
                        bold=False,
                        align=aligns[idx],
                        font_size=8,
                    )

            total_cells = table.add_row().cells
            total_values = [
                "",
                "",
                "",
                "",
                self._summary_label(),
                fmt_qty(group.summary.opening_balance),
                fmt_qty(group.summary.inflow_production),
                fmt_qty(group.summary.inflow_transfers_in),
                fmt_qty(group.summary.outflow_dispatch),
                fmt_qty(group.summary.outflow_transfers_out),
                safe_text(group.summary.total_boxes),
                safe_text(group.summary.total_pieces),
                fmt_qty(group.summary.closing_balance),
                "",
                "",
            ]
            total_aligns = [
                "center", "center", "left", "left", "left",
                "right", "right", "right", "right", "right",
                "right", "right", "right", "left", "center",
            ]
            for idx, value in enumerate(total_values):
                total_cells[idx].width = Inches(product_widths[idx])
                self._write_docx_cell(
                    total_cells[idx],
                    value,
                    bold=True,
                    align=total_aligns[idx],
                    font_size=8,
                    shading="F3F4F6",
                )

            document.add_paragraph("")

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _render_consumable_report_docx(
        self,
        report: ConsumableInventoryReportResponse,
        *,
        prepared_by: str | None = None,
    ) -> bytes:
        try:
            from docx import Document
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches, Pt
        except ImportError as exc:
            raise RuntimeError("python-docx is required to generate Word reports.") from exc

        prepared_by = self._resolve_prepared_by(prepared_by)
        document = Document()
        self._set_docx_landscape(document)

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("CONSUMABLE STORE INVENTORY REPORT")
        title_run.bold = True
        title_run.font.size = Pt(14)

        for line in [
            f"Prepared By: {prepared_by}",
            f"Report Type: {report.report_type.title()}",
            f"Period: {fmt_date(report.start_date)} to {fmt_date(report.end_date)}",
            f"Generated At: {report.generated_at.isoformat()}",
        ]:
            p = document.add_paragraph()
            r = p.add_run(line)
            r.font.size = Pt(9)

        consumable_widths = [
            0.40, 0.68, 0.80, 1.10, 1.35,
            0.55, 0.85, 0.85, 0.85, 1.25, 0.60,
        ]

        for group in report.groups:
            document.add_paragraph(group.label, style="Heading 2")

            headers = [
                "S.No.",
                "Date",
                "Store",
                "Item Category",
                "Item",
                "Unit",
                "Opening",
                "Issued Today",
                "Closing",
                "Remarks",
                "Checked",
            ]

            table = document.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            table.autofit = False
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for idx, header in enumerate(headers):
                cell = table.rows[0].cells[idx]
                cell.width = Inches(consumable_widths[idx])
                self._write_docx_cell(
                    cell,
                    header,
                    bold=True,
                    align="center",
                    font_size=8,
                    shading="E5E7EB",
                )

            for row in group.rows:
                cells = table.add_row().cells
                values = [
                    safe_text(row.serial_no),
                    fmt_date(row.entry_date),
                    row.store.value,
                    row.item_category_name,
                    row.item_name,
                    row.unit,
                    fmt_qty(row.opening_balance),
                    fmt_qty(row.issued_today),
                    fmt_qty(row.closing_balance),
                    safe_text(row.remarks),
                    safe_text(row.checked_by_initials),
                ]
                aligns = [
                    "center", "center", "left", "left", "left",
                    "center", "right", "right", "right", "left", "center",
                ]
                for idx, value in enumerate(values):
                    cells[idx].width = Inches(consumable_widths[idx])
                    self._write_docx_cell(
                        cells[idx],
                        value,
                        bold=False,
                        align=aligns[idx],
                        font_size=8,
                    )

            total_cells = table.add_row().cells
            total_values = [
                "",
                "",
                "",
                "",
                self._summary_label(),
                "",
                fmt_qty(group.summary.opening_balance),
                fmt_qty(group.summary.issued_today),
                fmt_qty(group.summary.closing_balance),
                "",
                "",
            ]
            total_aligns = [
                "center", "center", "left", "left", "left",
                "center", "right", "right", "right", "left", "center",
            ]
            for idx, value in enumerate(total_values):
                total_cells[idx].width = Inches(consumable_widths[idx])
                self._write_docx_cell(
                    total_cells[idx],
                    value,
                    bold=True,
                    align=total_aligns[idx],
                    font_size=8,
                    shading="F3F4F6",
                )

            document.add_paragraph("")

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _write_docx_cell(
        self,
        cell,
        text: str,
        *,
        bold: bool = False,
        align: str = "left",
        font_size: float = 8,
        shading: str | None = None,
    ) -> None:
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt

        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0

        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        run = p.add_run(safe_text(text))
        run.bold = bold
        run.font.size = Pt(font_size)

        if shading:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), shading)
            tc_pr.append(shd)

    def _set_docx_landscape(self, document) -> None:
        from docx.enum.section import WD_ORIENT
        from docx.shared import Inches

        section = document.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.left_margin = Inches(0.35)
        section.right_margin = Inches(0.35)
        section.top_margin = Inches(0.40)
        section.bottom_margin = Inches(0.40)

    # =========================================================================
    # OPTIONAL FILE INFO HELPERS
    # =========================================================================

    def export_product_file_info(
        self,
        payload: ProductInventoryReportRequest,
        *,
        prepared_by: str | None = None,
    ) -> InventoryExportFileResponse:
        file = self.export_product_inventory_report(payload, prepared_by=prepared_by)
        return file.to_schema(payload.report_type, payload.export_format or "pdf")

    def export_consumable_file_info(
        self,
        payload: ConsumableInventoryReportRequest,
        *,
        prepared_by: str | None = None,
    ) -> InventoryExportFileResponse:
        file = self.export_consumable_inventory_report(payload, prepared_by=prepared_by)
        return file.to_schema(payload.report_type, payload.export_format or "pdf")